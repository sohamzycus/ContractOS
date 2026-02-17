"""Word document (.docx) parser — extracts paragraphs, tables, headings with character offsets.

Includes bold-as-heading fallback for documents that don't use Word's
built-in Heading styles (common in legal contracts).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from contractos.tools.parsers import (
    ParsedDocument,
    ParsedParagraph,
    ParsedTable,
    ParsedTableCell,
)


def parse_docx(file_path: str | Path) -> ParsedDocument:
    """Parse a .docx file into a structured ParsedDocument.

    Extracts paragraphs with heading levels, tables with cell metadata,
    and computes character offsets into a concatenated full_text.

    If no built-in Heading styles are found, applies a bold-text
    fallback to detect section headings.
    """
    doc = Document(str(file_path))

    paragraphs: list[ParsedParagraph] = []
    tables: list[ParsedTable] = []
    full_text_parts: list[str] = []
    offset = 0
    para_idx = 0
    table_idx = 0

    # Walk the document body in order — paragraphs and tables interleaved
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            text = element.text or ""
            # Collect all runs
            runs = element.findall(qn("w:r"))
            text = ""
            for run in runs:
                t_elem = run.find(qn("w:t"))
                if t_elem is not None and t_elem.text:
                    text += t_elem.text
            if not text:
                # Try direct text content
                text = "".join(
                    t.text for t in element.iter(qn("w:t")) if t.text
                )

            heading_level = _get_heading_level(element)
            structural_path = f"body > p[{para_idx}]"
            if heading_level is not None:
                structural_path = f"body > heading[{heading_level}][{para_idx}]"

            char_start = offset
            char_end = offset + len(text)

            paragraphs.append(ParsedParagraph(
                text=text,
                char_start=char_start,
                char_end=char_end,
                structural_path=structural_path,
                heading_level=heading_level,
            ))

            if text:
                full_text_parts.append(text)
                offset = char_end + 1  # +1 for newline separator
            else:
                offset = char_end

            para_idx += 1

        elif tag == "tbl":
            table_cells: list[ParsedTableCell] = []
            rows = element.findall(qn("w:tr"))
            row_count = len(rows)
            col_count = 0
            table_char_start = offset

            for row_idx, row in enumerate(rows):
                cells = row.findall(qn("w:tc"))
                col_count = max(col_count, len(cells))
                for col_idx, cell in enumerate(cells):
                    cell_text = "".join(
                        t.text for t in cell.iter(qn("w:t")) if t.text
                    )
                    cell_char_start = offset
                    cell_char_end = offset + len(cell_text)

                    table_cells.append(ParsedTableCell(
                        text=cell_text,
                        row=row_idx,
                        col=col_idx,
                        char_start=cell_char_start,
                        char_end=cell_char_end,
                        structural_path=f"body > table[{table_idx}] > row[{row_idx}] > cell[{col_idx}]",
                    ))

                    if cell_text:
                        full_text_parts.append(cell_text)
                        offset = cell_char_end + 1
                    else:
                        offset = cell_char_end

            table_char_end = offset
            tables.append(ParsedTable(
                cells=table_cells,
                row_count=row_count,
                col_count=col_count,
                char_start=table_char_start,
                char_end=table_char_end,
                structural_path=f"body > table[{table_idx}]",
            ))
            table_idx += 1

    full_text = "\n".join(full_text_parts)
    word_count = len(full_text.split())

    # Bold-as-heading fallback: if no headings detected via styles,
    # re-scan for bold paragraphs that look like section headings
    has_headings = any(p.heading_level is not None for p in paragraphs)
    if not has_headings:
        paragraphs = _apply_bold_heading_fallback(paragraphs, doc)

    return ParsedDocument(
        paragraphs=paragraphs,
        tables=tables,
        full_text=full_text,
        page_count=0,  # .docx doesn't have reliable page count without rendering
        word_count=word_count,
    )


def _get_heading_level(element) -> int | None:
    """Extract heading level from a paragraph element, or None if not a heading."""
    pPr = element.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    style_val = pStyle.get(qn("w:val"), "")
    if style_val == "Title":
        return 0
    if style_val.startswith("Heading"):
        try:
            return int(style_val.replace("Heading", "").strip())
        except ValueError:
            return None
    return None


def _apply_bold_heading_fallback(
    paragraphs: list[ParsedParagraph],
    doc: Document,
) -> list[ParsedParagraph]:
    """Re-scan paragraphs for bold formatting and promote to headings.

    Only activates when the document has zero Heading-style paragraphs.
    Detects bold paragraphs that are short (< 80 chars) and look like
    section headings.
    """
    updated: list[ParsedParagraph] = []
    doc_paragraphs = list(doc.paragraphs)

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            updated.append(para)
            continue

        # Find the matching python-docx paragraph
        is_bold = False
        if i < len(doc_paragraphs):
            dp = doc_paragraphs[i]
            is_bold = _is_paragraph_bold(dp)

        heading_level = None
        if is_bold and _looks_like_heading(text):
            heading_level = 1

        # Also detect ALL-CAPS short text as headings
        if heading_level is None and _is_allcaps_heading(text):
            heading_level = 1

        # Also detect numbered section headings: "1. DEFINITIONS"
        if heading_level is None and _is_numbered_heading(text):
            heading_level = 1

        if heading_level is not None:
            updated.append(ParsedParagraph(
                text=para.text,
                char_start=para.char_start,
                char_end=para.char_end,
                structural_path=f"body > heading[{heading_level}][{i}]",
                heading_level=heading_level,
                page_number=para.page_number,
            ))
        else:
            updated.append(para)

    return updated


def _is_paragraph_bold(dp) -> bool:
    """Check if a python-docx Paragraph is bold (all runs or paragraph-level)."""
    if dp.paragraph_format and dp.style and dp.style.font and dp.style.font.bold:
        return True

    runs = dp.runs
    if not runs:
        return False

    # All non-empty runs must be bold
    non_empty_runs = [r for r in runs if r.text.strip()]
    if not non_empty_runs:
        return False

    return all(r.bold is True for r in non_empty_runs)


def _looks_like_heading(text: str) -> bool:
    """Check if text looks like a section heading (short, no trailing sentence)."""
    text = text.strip()
    if not text:
        return False
    if len(text) > 80:
        return False
    word_count = len(text.split())
    if word_count > 10:
        return False
    # Headings typically don't end with a period (unless it's a numbered heading like "1.")
    if text.endswith(".") and not text[-2].isdigit():
        return False
    return True


def _is_allcaps_heading(text: str) -> bool:
    """Check if text is an ALL-CAPS heading."""
    text = text.strip()
    if not text:
        return False
    if not (4 <= len(text) <= 60):
        return False
    word_count = len(text.split())
    if word_count > 8:
        return False
    if not any(c.isalpha() for c in text):
        return False
    alpha_chars = [c for c in text if c.isalpha()]
    return all(c.isupper() for c in alpha_chars)


def _is_numbered_heading(text: str) -> bool:
    """Check if text is a numbered section heading like '1. DEFINITIONS'."""
    import re
    text = text.strip()
    if not text:
        return False
    if len(text) > 80:
        return False
    # Match: "1.", "1.1", "12.", "12.1" followed by text starting with uppercase
    return bool(re.match(r'^\d+(?:\.\d+)*\.?\s+[A-Z]', text))
