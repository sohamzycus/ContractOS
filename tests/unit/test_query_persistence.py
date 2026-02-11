"""Tests for query persistence and multi-document support."""

from __future__ import annotations

import pytest
from httpx import ASGITransport, AsyncClient

from contractos.api.app import create_app
from contractos.api.deps import init_state, shutdown_state
from contractos.config import ContractOSConfig, LLMConfig, StorageConfig


@pytest.fixture
def test_config() -> ContractOSConfig:
    return ContractOSConfig(
        llm=LLMConfig(provider="mock"),
        storage=StorageConfig(path=":memory:"),
    )


@pytest.fixture
async def client(test_config: ContractOSConfig):
    init_state(test_config)
    app = create_app(test_config)
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    shutdown_state()


class TestQueryPersistence:
    """Test that queries are persisted as reasoning sessions."""

    async def test_ask_returns_session_id(self, client: AsyncClient) -> None:
        """Every Q&A response should include a session_id."""
        # Upload a document first
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Test Agreement", 0)
        doc.add_paragraph("This is a test contract between Alpha Corp and Beta Inc.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = await client.post(
            "/contracts/upload",
            files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        assert resp.status_code == 201
        doc_id = resp.json()["document_id"]

        # Ask a question
        resp = await client.post(
            "/query/ask",
            json={"question": "Who are the parties?", "document_id": doc_id},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert "session_id" in data
        assert data["session_id"].startswith("s-")

    async def test_chat_history_returns_sessions(self, client: AsyncClient) -> None:
        """Chat history should return previously asked questions."""
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("NDA Agreement", 0)
        doc.add_paragraph("Confidential information shall not be disclosed.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = await client.post(
            "/contracts/upload",
            files={"file": ("nda.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id = resp.json()["document_id"]

        # Ask two questions
        await client.post(
            "/query/ask",
            json={"question": "What is confidential?", "document_id": doc_id},
        )
        await client.post(
            "/query/ask",
            json={"question": "Who are the parties?", "document_id": doc_id},
        )

        # Get history
        resp = await client.get("/query/history")
        assert resp.status_code == 200
        history = resp.json()
        assert len(history) >= 2
        questions = [h["question"] for h in history]
        assert "What is confidential?" in questions
        assert "Who are the parties?" in questions

    async def test_chat_history_includes_answers(self, client: AsyncClient) -> None:
        """History items should include the answer."""
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Service Agreement", 0)
        doc.add_paragraph("The service fee is $10,000 per month.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = await client.post(
            "/contracts/upload",
            files={"file": ("svc.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id = resp.json()["document_id"]

        await client.post(
            "/query/ask",
            json={"question": "What is the fee?", "document_id": doc_id},
        )

        resp = await client.get("/query/history")
        history = resp.json()
        assert len(history) >= 1
        latest = history[0]
        assert latest["answer"] is not None
        assert latest["status"] == "completed"
        assert latest["document_ids"] is not None


class TestMultiDocumentQuery:
    """Test querying across multiple documents."""

    async def test_multi_doc_query_with_document_ids(self, client: AsyncClient) -> None:
        """Should accept document_ids list and query across both."""
        import io
        from docx import Document as DocxDocument

        # Upload doc 1
        doc1 = DocxDocument()
        doc1.add_heading("Vendor Agreement", 0)
        doc1.add_paragraph("The vendor shall deliver goods within 30 days.")
        buf1 = io.BytesIO()
        doc1.save(buf1)
        buf1.seek(0)

        resp1 = await client.post(
            "/contracts/upload",
            files={"file": ("vendor.docx", buf1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id_1 = resp1.json()["document_id"]

        # Upload doc 2
        doc2 = DocxDocument()
        doc2.add_heading("Service Agreement", 0)
        doc2.add_paragraph("The service provider guarantees 99.9% uptime.")
        buf2 = io.BytesIO()
        doc2.save(buf2)
        buf2.seek(0)

        resp2 = await client.post(
            "/contracts/upload",
            files={"file": ("service.docx", buf2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id_2 = resp2.json()["document_id"]

        # Query across both
        resp = await client.post(
            "/query/ask",
            json={
                "question": "What are the delivery and uptime guarantees?",
                "document_ids": [doc_id_1, doc_id_2],
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["answer"]
        assert len(data["document_ids"]) == 2

    async def test_single_doc_backward_compat(self, client: AsyncClient) -> None:
        """document_id (singular) should still work."""
        import io
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("Test", 0)
        doc.add_paragraph("Payment terms are Net 30.")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        resp = await client.post(
            "/contracts/upload",
            files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id = resp.json()["document_id"]

        resp = await client.post(
            "/query/ask",
            json={"question": "What are the payment terms?", "document_id": doc_id},
        )
        assert resp.status_code == 200
        assert len(resp.json()["document_ids"]) == 1

    async def test_missing_document_returns_404(self, client: AsyncClient) -> None:
        """Querying a non-existent document should return 404."""
        resp = await client.post(
            "/query/ask",
            json={"question": "test", "document_id": "doc-nonexistent"},
        )
        assert resp.status_code == 404

    async def test_no_document_returns_400(self, client: AsyncClient) -> None:
        """No document_id or document_ids should return 400."""
        resp = await client.post(
            "/query/ask",
            json={"question": "test"},
        )
        assert resp.status_code == 400
