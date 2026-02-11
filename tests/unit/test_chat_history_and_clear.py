"""Tests for chat history visibility, clear operations, and multi-doc with list."""

from __future__ import annotations

import io

import pytest
from docx import Document as DocxDocument
from httpx import ASGITransport, AsyncClient

from contractos.api.app import create_app
from contractos.api.deps import init_state, shutdown_state
from contractos.config import ContractOSConfig, LLMConfig, StorageConfig


@pytest.fixture
def test_config() -> ContractOSConfig:
    return ContractOSConfig(
        llm=LLMConfig(provider="mock"),
        storage=StorageConfig(path=":memory:"),
    )


@pytest.fixture
async def client(test_config: ContractOSConfig):
    init_state(test_config)
    app = create_app(test_config)
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    shutdown_state()


def _make_docx(title: str, body: str) -> io.BytesIO:
    doc = DocxDocument()
    doc.add_heading(title, 0)
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class TestChatHistoryVisibility:
    """User should be able to view chat history with full Q&A details."""

    async def test_history_empty_initially(self, client: AsyncClient) -> None:
        resp = await client.get("/query/history")
        assert resp.status_code == 200
        assert resp.json() == []

    async def test_history_shows_question_and_answer(self, client: AsyncClient) -> None:
        """After asking, history should contain question + answer."""
        buf = _make_docx("Test Agreement", "The payment term is Net 60 days.")
        resp = await client.post(
            "/contracts/upload",
            files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id = resp.json()["document_id"]

        await client.post(
            "/query/ask",
            json={"question": "What are the payment terms?", "document_id": doc_id},
        )

        resp = await client.get("/query/history")
        history = resp.json()
        assert len(history) == 1
        item = history[0]
        assert item["question"] == "What are the payment terms?"
        assert item["answer"] is not None
        assert item["status"] == "completed"
        assert item["session_id"].startswith("s-")
        assert doc_id in item["document_ids"]

    async def test_history_ordered_most_recent_first(self, client: AsyncClient) -> None:
        buf = _make_docx("NDA", "Confidential information must be protected.")
        resp = await client.post(
            "/contracts/upload",
            files={"file": ("nda.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id = resp.json()["document_id"]

        await client.post(
            "/query/ask", json={"question": "First question", "document_id": doc_id},
        )
        await client.post(
            "/query/ask", json={"question": "Second question", "document_id": doc_id},
        )

        resp = await client.get("/query/history")
        history = resp.json()
        assert len(history) == 2
        # Most recent first
        assert history[0]["question"] == "Second question"
        assert history[1]["question"] == "First question"


class TestClearOperations:
    """User should be able to clear uploaded files and history."""

    async def test_clear_chat_history(self, client: AsyncClient) -> None:
        """DELETE /query/history should clear all sessions."""
        buf = _make_docx("Test", "Some contract text here.")
        resp = await client.post(
            "/contracts/upload",
            files={"file": ("test.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id = resp.json()["document_id"]

        await client.post(
            "/query/ask", json={"question": "Test?", "document_id": doc_id},
        )

        # Verify history exists
        resp = await client.get("/query/history")
        assert len(resp.json()) == 1

        # Clear
        resp = await client.delete("/query/history")
        assert resp.status_code == 200
        data = resp.json()
        assert data["cleared"] == 1
        assert "Cleared" in data["message"]

        # Verify empty
        resp = await client.get("/query/history")
        assert resp.json() == []

    async def test_clear_all_contracts(self, client: AsyncClient) -> None:
        """DELETE /contracts/clear should remove all contracts and data."""
        buf1 = _make_docx("Contract A", "Payment is $10,000 per month.")
        buf2 = _make_docx("Contract B", "Termination requires 30 days notice.")

        resp1 = await client.post(
            "/contracts/upload",
            files={"file": ("a.docx", buf1, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        resp2 = await client.post(
            "/contracts/upload",
            files={"file": ("b.docx", buf2, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )
        doc_id_1 = resp1.json()["document_id"]
        doc_id_2 = resp2.json()["document_id"]

        # Verify contracts exist
        resp = await client.get("/contracts")
        assert len(resp.json()) == 2

        # Clear all
        resp = await client.delete("/contracts/clear")
        assert resp.status_code == 200
        data = resp.json()
        assert data["cleared_contracts"] == 2

        # Verify empty
        resp = await client.get("/contracts")
        assert resp.json() == []

        # Verify individual contracts are gone
        resp = await client.get(f"/contracts/{doc_id_1}")
        assert resp.status_code == 404

    async def test_clear_empty_is_safe(self, client: AsyncClient) -> None:
        """Clearing when nothing exists should not error."""
        resp = await client.delete("/contracts/clear")
        assert resp.status_code == 200
        assert resp.json()["cleared_contracts"] == 0

        resp = await client.delete("/query/history")
        assert resp.status_code == 200
        assert resp.json()["cleared"] == 0


class TestListContracts:
    """GET /contracts should list all uploaded documents."""

    async def test_list_empty(self, client: AsyncClient) -> None:
        resp = await client.get("/contracts")
        assert resp.status_code == 200
        assert resp.json() == []

    async def test_list_after_upload(self, client: AsyncClient) -> None:
        buf = _make_docx("Vendor Agreement", "The vendor shall deliver goods.")
        await client.post(
            "/contracts/upload",
            files={"file": ("vendor.docx", buf, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        )

        resp = await client.get("/contracts")
        contracts = resp.json()
        assert len(contracts) == 1
        assert contracts[0]["title"] == "vendor"
        assert contracts[0]["status"] == "indexed"
