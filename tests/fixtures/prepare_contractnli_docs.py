"""Extract real NDA documents from ContractNLI dataset and convert to DOCX.

Reads the 50 unique document names referenced in our legalbench train.tsv files,
finds the matching raw files from the ContractNLI dataset (607 NDAs), and converts
them to DOCX format for use with ContractOS.

Source: https://stanfordnlp.github.io/contract-nli/ (CC BY 4.0)
"""

from __future__ import annotations

import csv
import glob
import json
import os
import re
import shutil
from pathlib import Path
from urllib.parse import unquote

from docx import Document

FIXTURES_DIR = Path(__file__).parent
RAW_DIR = FIXTURES_DIR / "contractnli_raw" / "contract-nli" / "raw"
OUTPUT_DIR = FIXTURES_DIR / "contractnli_docs"
LEGALBENCH_DIR = FIXTURES_DIR / "legalbench"


def get_referenced_documents() -> set[str]:
    """Extract all unique document_name values from legalbench train.tsv files."""
    docs = set()
    for tsv in sorted(glob.glob(str(LEGALBENCH_DIR / "*/train.tsv"))):
        with open(tsv) as f:
            reader = csv.DictReader(f, delimiter="\t")
            for row in reader:
                name = row.get("document_name", "").strip()
                if name:
                    docs.add(name)
    return docs


def find_raw_file(doc_name: str) -> Path | None:
    """Find the raw file matching a document name (handles URL encoding)."""
    # Try exact match first
    path = RAW_DIR / doc_name
    if path.exists():
        return path

    # Try URL-decoded version
    decoded = unquote(doc_name)
    path = RAW_DIR / decoded
    if path.exists():
        return path

    # Try without leading space
    stripped = doc_name.strip()
    path = RAW_DIR / stripped
    if path.exists():
        return path

    return None


def convert_text_to_docx(text: str, title: str, output_path: Path) -> None:
    """Convert plain text content to a DOCX file."""
    doc = Document()
    doc.add_heading(title, 0)

    # Split into paragraphs and add them
    for para in text.split("\n"):
        para = para.strip()
        if para:
            # Detect headings (all caps or numbered sections)
            if (
                para.isupper()
                and len(para) < 100
                and not para.startswith("(")
            ):
                doc.add_heading(para, level=1)
            elif re.match(r"^\d+\.\s+[A-Z]", para):
                doc.add_heading(para, level=2)
            else:
                doc.add_paragraph(para)

    doc.save(str(output_path))


def extract_text_from_raw(raw_path: Path) -> str:
    """Extract text from a raw file (PDF, TXT, or HTML)."""
    suffix = raw_path.suffix.lower()

    if suffix == ".txt":
        return raw_path.read_text(errors="replace")

    if suffix == ".htm" or suffix == ".html":
        # Strip HTML tags for a rough text extraction
        html = raw_path.read_text(errors="replace")
        # Remove script/style blocks
        html = re.sub(r"<(script|style)[^>]*>.*?</\1>", "", html, flags=re.DOTALL | re.IGNORECASE)
        # Remove tags
        text = re.sub(r"<[^>]+>", " ", html)
        # Clean whitespace
        text = re.sub(r"\s+", " ", text).strip()
        # Re-add paragraph breaks at common boundaries
        text = re.sub(r"\s*\.\s+(\d+\.)", r".\n\n\1", text)
        return text

    if suffix == ".pdf":
        # Use the ContractNLI JSON to get the text (more reliable than PDF parsing)
        return ""  # Will be handled by JSON lookup

    return ""


def load_contractnli_texts() -> dict[str, str]:
    """Load full document texts from ContractNLI JSON files."""
    texts: dict[str, str] = {}
    json_dir = FIXTURES_DIR / "contractnli_raw" / "contract-nli"

    for json_file in ["train.json", "dev.json", "test.json"]:
        path = json_dir / json_file
        if not path.exists():
            continue
        with open(path) as f:
            data = json.load(f)
        for doc in data.get("documents", []):
            fname = doc.get("file_name", "")
            text = doc.get("text", "")
            if fname and text:
                texts[fname] = text
                # Also store stripped version for lookup
                texts[fname.strip()] = text

    return texts


def main() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)

    # Get documents referenced in our legalbench train.tsv files
    referenced = get_referenced_documents()
    print(f"Found {len(referenced)} unique documents referenced in legalbench train.tsv files")

    # Load full texts from ContractNLI JSON
    print("Loading ContractNLI JSON texts...")
    nli_texts = load_contractnli_texts()
    print(f"  Loaded texts for {len(nli_texts)} documents from JSON")

    converted = 0
    skipped = 0

    for doc_name in sorted(referenced):
        # Clean the name for output
        clean_name = unquote(doc_name).replace(" ", "_").replace("%20", "_")
        clean_name = re.sub(r"[^\w\-.]", "_", clean_name)
        docx_name = re.sub(r"\.(pdf|txt|htm|html)$", ".docx", clean_name, flags=re.IGNORECASE)
        output_path = OUTPUT_DIR / docx_name

        # Try to get text from JSON first (most reliable)
        text = nli_texts.get(doc_name, "")

        if not text:
            # Try URL-decoded name
            decoded = unquote(doc_name)
            text = nli_texts.get(decoded, "")

        if not text:
            # Try to read from raw file
            raw_path = find_raw_file(doc_name)
            if raw_path:
                text = extract_text_from_raw(raw_path)

        if not text or len(text) < 100:
            print(f"  SKIP: {doc_name} (no text found or too short)")
            skipped += 1
            continue

        # Generate a title from the filename
        title = re.sub(r"\.(pdf|txt|htm|html)$", "", unquote(doc_name), flags=re.IGNORECASE)
        title = title.replace("-", " ").replace("_", " ").replace("%20", " ").strip()
        if len(title) > 80:
            title = title[:80]

        convert_text_to_docx(text, title, output_path)
        converted += 1
        word_count = len(text.split())
        print(f"  OK: {docx_name} ({word_count} words)")

    print(f"\nDone: {converted} converted, {skipped} skipped")
    print(f"Output directory: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
