"""Generate complex real-world procurement contract fixtures for ContractOS.

These fixtures simulate production-grade contracts with:
- Multi-party agreements with complex entity aliases
- Nested clause structures with cross-references
- SLA tables with penalty tiers
- Price escalation schedules
- Indemnity caps and liability limitations
- Multi-location coverage with regional breakdowns
- Change order procedures
- Regulatory compliance requirements
- Insurance and bonding requirements
- Dispute resolution mechanisms

Run:
    python tests/fixtures/create_complex_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent


def create_it_outsourcing_agreement_docx() -> None:
    """Create a complex IT outsourcing agreement .docx.

    Simulates a real-world multi-year IT outsourcing contract between
    a large enterprise and a managed services provider, covering:
    - Infrastructure management across 12 locations
    - Application support for 8 enterprise systems
    - SLA framework with 5 severity levels and penalty tiers
    - Price escalation tied to CPI
    - Complex indemnity and liability structure
    - Data protection and GDPR compliance
    - Business continuity and disaster recovery
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title Page ──
    title = doc.add_heading("IT OUTSOURCING SERVICES AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph("Contract Reference: ITO-2025-0847")
    doc.add_paragraph("Effective Date: March 1, 2025")
    doc.add_paragraph("Expiry Date: February 28, 2030")
    doc.add_paragraph("Total Contract Value: $47,500,000.00")
    doc.add_paragraph("")

    # ── Section 1: Definitions and Interpretation ──
    doc.add_heading("1. Definitions and Interpretation", level=1)

    doc.add_heading("1.1 Parties", level=2)
    doc.add_paragraph(
        'This IT Outsourcing Services Agreement ("Agreement" or "Contract") '
        'is entered into as of March 1, 2025 ("Effective Date") by and between:'
    )
    doc.add_paragraph(
        'Meridian Global Holdings Inc., a corporation organized under the laws of '
        'the State of Delaware, with its principal office at 200 Park Avenue, '
        'New York, NY 10166 (hereinafter referred to as "Client" or "Meridian");'
    )
    doc.add_paragraph("and")
    doc.add_paragraph(
        'TechServe Solutions Pvt. Ltd., a company incorporated under the laws of '
        'India, with its registered office at Cyber Tower, HITEC City, Hyderabad, '
        'Telangana 500081 (hereinafter referred to as "Service Provider" or '
        '"TechServe" or "Vendor");'
    )
    doc.add_paragraph(
        "(collectively referred to as the \"Parties\" and individually as a \"Party\")."
    )

    doc.add_heading("1.2 Key Definitions", level=2)
    definitions = [
        ('"Authorized Representative"', 'means a person designated in writing by a Party to act on its behalf under this Agreement, as listed in Schedule F.'),
        ('"Business Day"', 'means any day other than a Saturday, Sunday, or public holiday in the jurisdiction of the affected Service Location.'),
        ('"Change Order"', 'means a written request to modify the scope, schedule, or pricing of Services, processed in accordance with Section 14.'),
        ('"Confidential Information"', 'means all information disclosed by either Party that is marked as confidential or that a reasonable person would understand to be confidential, including but not limited to trade secrets, business plans, financial data, customer lists, technical specifications, and source code, as further described in Section 9.'),
        ('"Critical System"', 'means any system classified as Severity 1 or Severity 2 in the SLA Framework set out in Schedule C.'),
        ('"Data Protection Laws"', 'means the General Data Protection Regulation (EU) 2016/679 ("GDPR"), the California Consumer Privacy Act ("CCPA"), the Indian Information Technology Act 2000, and any other applicable data protection legislation in the jurisdictions where Services are performed.'),
        ('"Deliverables"', 'means all work products, reports, documentation, software, and other materials created by the Service Provider in the performance of the Services.'),
        ('"Force Majeure Event"', 'has the meaning set forth in Section 16.'),
        ('"Intellectual Property" or "IP"', 'means all patents, copyrights, trademarks, trade secrets, know-how, and other intellectual property rights.'),
        ('"Key Personnel"', 'means the individuals identified in Schedule F who are essential to the delivery of Services and may not be replaced without the Client\'s prior written consent.'),
        ('"Service Credits"', 'means the financial credits payable by the Service Provider to the Client in the event of SLA breaches, calculated in accordance with Schedule C, Section 4.'),
        ('"Service Levels" or "SLAs"', 'means the performance standards and metrics set forth in Schedule C.'),
        ('"Transition Period"', 'means the period of one hundred and twenty (120) days from the Effective Date during which the Service Provider assumes responsibility for the Services from the Client\'s incumbent provider.'),
    ]
    for term, defn in definitions:
        doc.add_paragraph(f"{term} {defn}")

    # ── Section 2: Scope of Services ──
    doc.add_heading("2. Scope of Services", level=1)

    doc.add_heading("2.1 Infrastructure Management Services", level=2)
    doc.add_paragraph(
        "The Service Provider shall provide comprehensive infrastructure management "
        "services covering all hardware, software, and network components listed in "
        "Schedule A across all Service Locations identified in Schedule B. The "
        "infrastructure management services include, but are not limited to:"
    )
    items_infra = [
        "Server administration and monitoring (physical and virtual) for 2,400+ servers",
        "Network management including LAN, WAN, SD-WAN, and VPN infrastructure",
        "Storage management across SAN, NAS, and cloud storage tiers",
        "Database administration for Oracle, SQL Server, PostgreSQL, and MongoDB instances",
        "End-user computing support for 15,000 desktop and laptop devices",
        "Data center operations at Tier III and Tier IV facilities",
        "Cloud infrastructure management across AWS, Azure, and GCP environments",
        "Security operations center (SOC) monitoring on a 24x7x365 basis",
    ]
    for item in items_infra:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.2 Application Support Services", level=2)
    doc.add_paragraph(
        "The Service Provider shall provide application support and maintenance for "
        "the enterprise systems listed in Schedule A, Part II. Application support "
        "services shall include:"
    )
    items_app = [
        "Level 2 and Level 3 application support for SAP S/4HANA (Finance, HR, SCM, MM modules)",
        "Salesforce CRM administration and customization support",
        "Oracle E-Business Suite maintenance and patch management",
        "Custom application support for Meridian's proprietary trading platform (\"MeridianTrade\")",
        "ServiceNow ITSM platform administration and workflow development",
        "Microsoft 365 and Azure Active Directory management",
        "Business intelligence support for Tableau and Power BI environments",
        "Integration middleware support for MuleSoft and Apache Kafka platforms",
    ]
    for item in items_app:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2.3 Excluded Services", level=2)
    doc.add_paragraph(
        "The following services are expressly excluded from the scope of this Agreement "
        "and shall require a separate Change Order pursuant to Section 14 if the Client "
        "wishes to include them:"
    )
    excluded = [
        "New application development (greenfield projects)",
        "Hardware procurement and capital expenditure",
        "Telecommunications carrier management",
        "Physical security of Client premises",
        "End-user training programs",
    ]
    for item in excluded:
        doc.add_paragraph(item, style="List Bullet")

    # ── Section 3: Term and Renewal ──
    doc.add_heading("3. Term and Renewal", level=1)

    doc.add_heading("3.1 Initial Term", level=2)
    doc.add_paragraph(
        "This Agreement shall commence on the Effective Date and continue for an "
        "initial term of five (5) years (the \"Initial Term\"), expiring on "
        "February 28, 2030, unless earlier terminated in accordance with Section 12."
    )

    doc.add_heading("3.2 Renewal", level=2)
    doc.add_paragraph(
        "Upon expiration of the Initial Term, this Agreement shall automatically "
        "renew for successive periods of two (2) years each (each a \"Renewal Term\"), "
        "unless either Party provides written notice of non-renewal at least one "
        "hundred and eighty (180) days prior to the expiration of the then-current term."
    )

    doc.add_heading("3.3 Transition Period", level=2)
    doc.add_paragraph(
        "The Transition Period shall commence on the Effective Date and conclude "
        "no later than one hundred and twenty (120) days thereafter (the \"Transition "
        "Completion Date\"). During the Transition Period, the Service Provider shall:"
    )
    transition_items = [
        "Conduct a comprehensive assessment of the Client's existing IT environment",
        "Develop and execute a detailed transition plan approved by the Client",
        "Assume operational responsibility for all in-scope services in a phased manner",
        "Achieve steady-state operations by the Transition Completion Date",
        "Provide weekly transition status reports to the Client's Authorized Representative",
    ]
    for item in transition_items:
        doc.add_paragraph(item, style="List Bullet")

    # ── Section 4: Pricing and Payment ──
    doc.add_heading("4. Pricing and Payment", level=1)

    doc.add_heading("4.1 Contract Value", level=2)
    doc.add_paragraph(
        "The total contract value for the Initial Term is Forty-Seven Million Five "
        "Hundred Thousand US Dollars ($47,500,000.00), structured as follows:"
    )

    # Pricing table
    doc.add_heading("4.1.1 Annual Fee Schedule", level=3)
    pricing_table = doc.add_table(rows=7, cols=4)
    pricing_table.style = "Table Grid"
    headers = ["Year", "Base Fee", "Variable Component (est.)", "Total (est.)"]
    for i, h in enumerate(headers):
        pricing_table.rows[0].cells[i].text = h
    pricing_data = [
        ("Year 1 (2025-26)", "$7,200,000", "$1,800,000", "$9,000,000"),
        ("Year 2 (2026-27)", "$7,416,000", "$1,854,000", "$9,270,000"),
        ("Year 3 (2027-28)", "$7,638,480", "$1,909,620", "$9,548,100"),
        ("Year 4 (2028-29)", "$7,867,634", "$1,966,909", "$9,834,543"),
        ("Year 5 (2029-30)", "$8,103,663", "$2,025,916", "$10,129,579"),
        ("TOTAL", "$38,225,777", "$9,556,445", "$47,782,222"),
    ]
    for row_idx, (year, base, var, total) in enumerate(pricing_data, start=1):
        pricing_table.rows[row_idx].cells[0].text = year
        pricing_table.rows[row_idx].cells[1].text = base
        pricing_table.rows[row_idx].cells[2].text = var
        pricing_table.rows[row_idx].cells[3].text = total

    doc.add_heading("4.2 Price Escalation", level=2)
    doc.add_paragraph(
        "The Base Fee shall be subject to an annual escalation of three percent (3%) "
        "or the Consumer Price Index (CPI) increase for the preceding calendar year, "
        "whichever is lower (the \"Escalation Cap\"). The escalation shall apply from "
        "Year 2 onwards and shall be calculated on the Base Fee of the immediately "
        "preceding year. The Service Provider shall provide the Client with a revised "
        "fee schedule no later than sixty (60) days before the start of each contract year."
    )

    doc.add_heading("4.3 Payment Terms", level=2)
    doc.add_paragraph(
        "The Client shall pay the Service Provider in equal monthly installments of "
        "the Base Fee, payable within forty-five (45) days of receipt of a valid "
        "invoice. Variable components shall be invoiced quarterly in arrears based on "
        "actual consumption, subject to the reconciliation process described in "
        "Section 4.5. All payments shall be made in US Dollars by wire transfer to "
        "the account specified in Schedule E."
    )

    doc.add_heading("4.4 Service Credits and Deductions", level=2)
    doc.add_paragraph(
        "Service Credits earned by the Client pursuant to Schedule C shall be applied "
        "as deductions against the next monthly invoice. The maximum aggregate Service "
        "Credits in any calendar month shall not exceed fifteen percent (15%) of the "
        "monthly Base Fee (the \"Monthly Credit Cap\"). Service Credits that exceed the "
        "Monthly Credit Cap shall be carried forward to subsequent months. Notwithstanding "
        "the foregoing, if aggregate Service Credits in any rolling twelve (12) month "
        "period exceed twenty percent (20%) of the annual Base Fee, the Client shall "
        "have the right to terminate this Agreement for cause pursuant to Section 12.2."
    )

    doc.add_heading("4.5 Quarterly Reconciliation", level=2)
    doc.add_paragraph(
        "Within thirty (30) days following the end of each calendar quarter, the "
        "Parties shall conduct a reconciliation of all variable charges, Service "
        "Credits, and any disputed amounts. Any net amount due to either Party shall "
        "be settled within fifteen (15) Business Days of the reconciliation completion."
    )

    # ── Section 5: Service Levels ──
    doc.add_heading("5. Service Levels and Performance Standards", level=1)

    doc.add_heading("5.1 SLA Framework", level=2)
    doc.add_paragraph(
        "The Service Provider shall meet or exceed the Service Levels set forth in "
        "Schedule C. The SLA framework is structured around five (5) severity levels, "
        "each with defined response times, resolution times, and associated Service "
        "Credits for non-compliance."
    )

    # SLA summary table
    sla_table = doc.add_table(rows=6, cols=5)
    sla_table.style = "Table Grid"
    sla_headers = ["Severity", "Description", "Response Time", "Resolution Time", "Service Credit (% monthly fee)"]
    for i, h in enumerate(sla_headers):
        sla_table.rows[0].cells[i].text = h
    sla_data = [
        ("Severity 1 (Critical)", "Complete system outage affecting >50% of users", "15 minutes", "4 hours", "5% per incident"),
        ("Severity 2 (High)", "Major degradation affecting business-critical functions", "30 minutes", "8 hours", "3% per incident"),
        ("Severity 3 (Medium)", "Partial service impact with workaround available", "2 hours", "24 hours", "1% per incident"),
        ("Severity 4 (Low)", "Minor issue with minimal business impact", "4 hours", "72 hours", "0.5% per incident"),
        ("Severity 5 (Informational)", "Service request or enhancement", "1 Business Day", "5 Business Days", "N/A"),
    ]
    for row_idx, (sev, desc, resp, resol, credit) in enumerate(sla_data, start=1):
        sla_table.rows[row_idx].cells[0].text = sev
        sla_table.rows[row_idx].cells[1].text = desc
        sla_table.rows[row_idx].cells[2].text = resp
        sla_table.rows[row_idx].cells[3].text = resol
        sla_table.rows[row_idx].cells[4].text = credit

    doc.add_heading("5.2 Availability Targets", level=2)
    avail_table = doc.add_table(rows=5, cols=3)
    avail_table.style = "Table Grid"
    avail_headers = ["Service Category", "Monthly Uptime Target", "Measurement Window"]
    for i, h in enumerate(avail_headers):
        avail_table.rows[0].cells[i].text = h
    avail_data = [
        ("Critical Production Systems", "99.99%", "24x7x365"),
        ("Business Applications (Tier 1)", "99.95%", "24x7x365"),
        ("Business Applications (Tier 2)", "99.9%", "Business Hours (8am-8pm local)"),
        ("Development/Test Environments", "99.5%", "Business Hours (8am-6pm local)"),
    ]
    for row_idx, (cat, target, window) in enumerate(avail_data, start=1):
        avail_table.rows[row_idx].cells[0].text = cat
        avail_table.rows[row_idx].cells[1].text = target
        avail_table.rows[row_idx].cells[2].text = window

    doc.add_heading("5.3 Performance Reporting", level=2)
    doc.add_paragraph(
        "The Service Provider shall deliver the following reports to the Client's "
        "Authorized Representative:"
    )
    reports = [
        "Daily: Incident summary and open ticket status (by 9:00 AM EST)",
        "Weekly: SLA performance dashboard with trend analysis",
        "Monthly: Comprehensive service report including uptime metrics, incident analysis, capacity planning, and Service Credit calculations",
        "Quarterly: Executive summary with strategic recommendations and continuous improvement initiatives",
        "Annually: Year-in-review report with benchmarking against industry standards",
    ]
    for item in reports:
        doc.add_paragraph(item, style="List Bullet")

    # ── Section 6: Governance ──
    doc.add_heading("6. Governance and Relationship Management", level=1)

    doc.add_heading("6.1 Governance Structure", level=2)
    doc.add_paragraph(
        "The Parties shall establish a multi-tiered governance structure as follows:"
    )
    governance = [
        "Strategic Level: Executive Steering Committee meeting quarterly, comprising C-level representatives from both Parties, responsible for strategic direction and dispute escalation",
        "Tactical Level: Service Delivery Board meeting monthly, comprising senior managers, responsible for service performance review and improvement planning",
        "Operational Level: Weekly operations meetings between operational leads, responsible for day-to-day service delivery and incident management",
    ]
    for item in governance:
        doc.add_paragraph(item, style="List Bullet")

    # ── Section 7: Personnel ──
    doc.add_heading("7. Personnel and Staffing", level=1)
    doc.add_paragraph(
        "The Service Provider shall maintain a dedicated team of no fewer than "
        "three hundred and fifty (350) full-time equivalent (FTE) personnel for the "
        "delivery of Services. Of these, at least forty (40) FTEs shall be based at "
        "the Client's primary locations as specified in Schedule B. Key Personnel "
        "identified in Schedule F shall not be reassigned or replaced without the "
        "Client's prior written consent, which shall not be unreasonably withheld. "
        "The Service Provider shall ensure that all personnel assigned to this "
        "engagement possess the certifications and qualifications specified in "
        "Schedule F, Section 3."
    )

    # ── Section 8: Intellectual Property ──
    doc.add_heading("8. Intellectual Property Rights", level=1)

    doc.add_heading("8.1 Client IP", level=2)
    doc.add_paragraph(
        "All Intellectual Property owned by the Client prior to the Effective Date "
        "or developed independently by the Client during the Term (\"Client IP\") "
        "shall remain the sole property of the Client. The Service Provider is "
        "granted a limited, non-exclusive, non-transferable license to use Client "
        "IP solely for the purpose of performing the Services during the Term."
    )

    doc.add_heading("8.2 Service Provider IP", level=2)
    doc.add_paragraph(
        "All Intellectual Property owned by the Service Provider prior to the "
        "Effective Date or developed independently by the Service Provider "
        "(\"Service Provider IP\") shall remain the property of the Service Provider. "
        "The Client is granted a perpetual, irrevocable, royalty-free license to use "
        "any Service Provider IP that is embedded in the Deliverables."
    )

    doc.add_heading("8.3 Jointly Developed IP", level=2)
    doc.add_paragraph(
        "Any Intellectual Property developed jointly by the Parties during the "
        "performance of this Agreement (\"Joint IP\") shall be jointly owned. "
        "Neither Party shall license or assign Joint IP to any third party without "
        "the prior written consent of the other Party. Disputes regarding the "
        "classification of IP shall be resolved through the governance process "
        "described in Section 6, escalating to the Executive Steering Committee "
        "if not resolved at the operational level within thirty (30) days."
    )

    # ── Section 9: Confidentiality ──
    doc.add_heading("9. Confidentiality and Data Protection", level=1)

    doc.add_heading("9.1 Confidentiality Obligations", level=2)
    doc.add_paragraph(
        "Each Party (the \"Receiving Party\") agrees to hold in strict confidence "
        "all Confidential Information received from the other Party (the \"Disclosing "
        "Party\") and shall not disclose such information to any third party without "
        "the prior written consent of the Disclosing Party, except as required by "
        "law or regulation. The confidentiality obligations under this Section shall "
        "survive the termination or expiration of this Agreement for a period of "
        "five (5) years."
    )

    doc.add_heading("9.2 Data Protection Compliance", level=2)
    doc.add_paragraph(
        "The Service Provider shall comply with all applicable Data Protection Laws "
        "in the processing of personal data on behalf of the Client. The Service "
        "Provider shall act as a data processor under GDPR and shall execute a "
        "Data Processing Agreement (\"DPA\") in the form attached as Schedule G. "
        "The Service Provider shall implement appropriate technical and organizational "
        "measures to ensure a level of security appropriate to the risk, including "
        "but not limited to encryption of data at rest (AES-256) and in transit "
        "(TLS 1.3), access controls, regular security assessments, and incident "
        "response procedures."
    )

    doc.add_heading("9.3 Data Breach Notification", level=2)
    doc.add_paragraph(
        "In the event of a data breach involving the Client's personal data, the "
        "Service Provider shall notify the Client within twenty-four (24) hours of "
        "becoming aware of the breach. The notification shall include the nature of "
        "the breach, categories of data affected, estimated number of data subjects "
        "affected, and the measures taken or proposed to address the breach. The "
        "Service Provider shall cooperate fully with the Client in investigating "
        "and remediating the breach, and shall bear all costs associated with the "
        "breach to the extent caused by the Service Provider's negligence or "
        "non-compliance with this Agreement."
    )

    # ── Section 10: Insurance ──
    doc.add_heading("10. Insurance", level=1)
    doc.add_paragraph(
        "The Service Provider shall maintain, at its own expense, the following "
        "insurance coverage throughout the Term and for a period of two (2) years "
        "following termination or expiration:"
    )
    insurance_table = doc.add_table(rows=6, cols=3)
    insurance_table.style = "Table Grid"
    ins_headers = ["Coverage Type", "Minimum Limit", "Deductible (max)"]
    for i, h in enumerate(ins_headers):
        insurance_table.rows[0].cells[i].text = h
    ins_data = [
        ("Commercial General Liability", "$10,000,000 per occurrence", "$100,000"),
        ("Professional Liability (E&O)", "$25,000,000 per claim", "$250,000"),
        ("Cyber Liability", "$20,000,000 per claim", "$500,000"),
        ("Workers' Compensation", "Statutory limits", "N/A"),
        ("Employer's Liability", "$5,000,000 per occurrence", "$50,000"),
    ]
    for row_idx, (cov, limit, ded) in enumerate(ins_data, start=1):
        insurance_table.rows[row_idx].cells[0].text = cov
        insurance_table.rows[row_idx].cells[1].text = limit
        insurance_table.rows[row_idx].cells[2].text = ded

    # ── Section 11: Indemnification and Liability ──
    doc.add_heading("11. Indemnification and Limitation of Liability", level=1)

    doc.add_heading("11.1 Service Provider Indemnification", level=2)
    doc.add_paragraph(
        "The Service Provider shall indemnify, defend, and hold harmless the Client, "
        "its affiliates, officers, directors, employees, and agents from and against "
        "all claims, damages, losses, liabilities, costs, and expenses (including "
        "reasonable attorneys' fees) arising out of or relating to: (a) the Service "
        "Provider's breach of this Agreement; (b) the Service Provider's negligence "
        "or willful misconduct; (c) any infringement of third-party Intellectual "
        "Property rights by the Deliverables; (d) any violation of applicable Data "
        "Protection Laws by the Service Provider; or (e) any personal injury or "
        "property damage caused by the Service Provider's personnel."
    )

    doc.add_heading("11.2 Limitation of Liability", level=2)
    doc.add_paragraph(
        "Except for claims arising from (i) breach of confidentiality obligations "
        "under Section 9, (ii) indemnification obligations under Section 11.1(c) "
        "and (d), or (iii) willful misconduct or gross negligence, the aggregate "
        "liability of either Party under this Agreement shall not exceed two hundred "
        "percent (200%) of the annual Base Fee for the then-current contract year "
        "(the \"Liability Cap\"). For the avoidance of doubt, the Liability Cap for "
        "Year 1 shall be Fourteen Million Four Hundred Thousand US Dollars "
        "($14,400,000.00)."
    )

    doc.add_heading("11.3 Exclusion of Consequential Damages", level=2)
    doc.add_paragraph(
        "Neither Party shall be liable to the other for any indirect, incidental, "
        "consequential, special, or punitive damages, including but not limited to "
        "loss of profits, loss of revenue, loss of data, or loss of business "
        "opportunity, regardless of the cause of action or the theory of liability, "
        "even if such Party has been advised of the possibility of such damages. "
        "This exclusion shall not apply to: (a) the Service Provider's indemnification "
        "obligations under Section 11.1; (b) breaches of confidentiality under "
        "Section 9; or (c) damages arising from willful misconduct."
    )

    # ── Section 12: Termination ──
    doc.add_heading("12. Termination", level=1)

    doc.add_heading("12.1 Termination for Convenience", level=2)
    doc.add_paragraph(
        "The Client may terminate this Agreement for convenience by providing the "
        "Service Provider with one hundred and eighty (180) days' prior written "
        "notice. In the event of termination for convenience, the Client shall pay "
        "the Service Provider: (a) all fees due for Services performed through the "
        "termination date; (b) a termination fee equal to six (6) months of the "
        "Base Fee; and (c) reasonable, documented wind-down costs not to exceed "
        "Five Hundred Thousand US Dollars ($500,000.00)."
    )

    doc.add_heading("12.2 Termination for Cause", level=2)
    doc.add_paragraph(
        "Either Party may terminate this Agreement for cause upon the occurrence "
        "of any of the following events: (a) material breach that remains uncured "
        "for thirty (30) days after written notice; (b) insolvency, bankruptcy, or "
        "appointment of a receiver; (c) failure to meet Severity 1 SLA targets for "
        "three (3) consecutive months; (d) aggregate Service Credits exceeding "
        "twenty percent (20%) of the annual Base Fee in any rolling twelve (12) "
        "month period, as referenced in Section 4.4; or (e) a data breach caused "
        "by the Service Provider's gross negligence, as described in Section 9.3."
    )

    doc.add_heading("12.3 Termination Assistance", level=2)
    doc.add_paragraph(
        "Upon termination or expiration of this Agreement, the Service Provider "
        "shall provide termination assistance services for a period of up to twelve "
        "(12) months (the \"Termination Assistance Period\") to facilitate the "
        "orderly transition of Services to the Client or its designated successor. "
        "Termination assistance shall be provided at the rates specified in Schedule "
        "E, Section 3. The Service Provider shall cooperate fully and in good faith "
        "during the transition, including providing access to all documentation, "
        "knowledge bases, and operational procedures."
    )

    # ── Section 13: Business Continuity ──
    doc.add_heading("13. Business Continuity and Disaster Recovery", level=1)
    doc.add_paragraph(
        "The Service Provider shall maintain a comprehensive Business Continuity "
        "Plan (\"BCP\") and Disaster Recovery Plan (\"DRP\") that ensure the "
        "continuity of Services in the event of a disruption. The BCP/DRP shall "
        "include: (a) Recovery Time Objective (RTO) of four (4) hours for Critical "
        "Systems and twenty-four (24) hours for non-critical systems; (b) Recovery "
        "Point Objective (RPO) of one (1) hour for Critical Systems and four (4) "
        "hours for non-critical systems; (c) annual testing of the BCP/DRP with "
        "the Client's participation; and (d) a geographically diverse disaster "
        "recovery site located at least 200 kilometers from the primary data center. "
        "The Service Provider shall provide the Client with a copy of the BCP/DRP "
        "within sixty (60) days of the Effective Date and shall update it annually "
        "or upon any material change to the service delivery environment."
    )

    # ── Section 14: Change Management ──
    doc.add_heading("14. Change Management", level=1)

    doc.add_heading("14.1 Change Order Process", level=2)
    doc.add_paragraph(
        "Either Party may request a change to the scope, schedule, or pricing of "
        "Services by submitting a Change Order in the form set out in Schedule H. "
        "The receiving Party shall respond to a Change Order within ten (10) Business "
        "Days with either an acceptance, rejection, or counter-proposal. No change "
        "shall be effective until a Change Order is signed by the Authorized "
        "Representatives of both Parties."
    )

    doc.add_heading("14.2 Emergency Changes", level=2)
    doc.add_paragraph(
        "In the event of an emergency requiring immediate changes to prevent or "
        "mitigate a Severity 1 incident, the Service Provider may implement changes "
        "without prior written approval, provided that: (a) the Service Provider "
        "notifies the Client's Authorized Representative within two (2) hours of "
        "implementing the change; (b) a retrospective Change Order is submitted "
        "within five (5) Business Days; and (c) the change is documented in the "
        "incident report."
    )

    # ── Section 15: Compliance ──
    doc.add_heading("15. Regulatory Compliance", level=1)
    doc.add_paragraph(
        "The Service Provider shall comply with all applicable laws, regulations, "
        "and industry standards in the performance of Services, including but not "
        "limited to: SOC 2 Type II certification, ISO 27001 certification, PCI DSS "
        "compliance (for payment-related systems), HIPAA compliance (for healthcare "
        "data), and any sector-specific regulations applicable to the Client's "
        "industry. The Service Provider shall provide evidence of compliance upon "
        "request and shall notify the Client within five (5) Business Days of any "
        "change in its compliance status."
    )

    # ── Section 16: Force Majeure ──
    doc.add_heading("16. Force Majeure", level=1)
    doc.add_paragraph(
        "Neither Party shall be liable for any failure or delay in performing its "
        "obligations under this Agreement to the extent that such failure or delay "
        "is caused by a Force Majeure Event. A \"Force Majeure Event\" means any "
        "event beyond the reasonable control of the affected Party, including but "
        "not limited to: natural disasters, epidemics, pandemics, acts of war or "
        "terrorism, government actions, labor disputes (other than those involving "
        "the affected Party's own employees), or failure of third-party "
        "telecommunications or power infrastructure. The affected Party shall: "
        "(a) notify the other Party within forty-eight (48) hours of the occurrence "
        "of the Force Majeure Event; (b) use commercially reasonable efforts to "
        "mitigate the impact; and (c) resume performance as soon as reasonably "
        "practicable. If a Force Majeure Event continues for more than ninety (90) "
        "consecutive days, either Party may terminate this Agreement upon thirty "
        "(30) days' written notice."
    )

    # ── Section 17: Dispute Resolution ──
    doc.add_heading("17. Dispute Resolution", level=1)

    doc.add_heading("17.1 Escalation", level=2)
    doc.add_paragraph(
        "Any dispute arising out of or in connection with this Agreement shall first "
        "be referred to the operational leads of both Parties for resolution within "
        "fifteen (15) Business Days. If unresolved, the dispute shall be escalated "
        "to the Executive Steering Committee for resolution within thirty (30) "
        "Business Days."
    )

    doc.add_heading("17.2 Mediation", level=2)
    doc.add_paragraph(
        "If the dispute is not resolved through the escalation process described in "
        "Section 17.1, the Parties shall submit the dispute to mediation administered "
        "by the International Chamber of Commerce (\"ICC\") in accordance with its "
        "mediation rules. The mediation shall take place in New York, New York."
    )

    doc.add_heading("17.3 Arbitration", level=2)
    doc.add_paragraph(
        "If the dispute is not resolved through mediation within sixty (60) days, "
        "the dispute shall be finally resolved by arbitration administered by the "
        "ICC under its Rules of Arbitration. The arbitration shall be conducted by "
        "three (3) arbitrators, with each Party appointing one arbitrator and the "
        "two appointed arbitrators selecting the third. The seat of arbitration "
        "shall be New York, New York, and the language of the arbitration shall be "
        "English. The arbitral award shall be final and binding on both Parties."
    )

    # ── Section 18: General Provisions ──
    doc.add_heading("18. General Provisions", level=1)

    doc.add_heading("18.1 Governing Law", level=2)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with the "
        "laws of the State of New York, without regard to its conflict of laws principles."
    )

    doc.add_heading("18.2 Entire Agreement", level=2)
    doc.add_paragraph(
        "This Agreement, together with all Schedules and Exhibits attached hereto, "
        "constitutes the entire agreement between the Parties with respect to the "
        "subject matter hereof and supersedes all prior negotiations, representations, "
        "warranties, and agreements between the Parties."
    )

    doc.add_heading("18.3 Amendments", level=2)
    doc.add_paragraph(
        "No amendment or modification of this Agreement shall be effective unless "
        "made in writing and signed by the Authorized Representatives of both Parties."
    )

    doc.add_heading("18.4 Assignment", level=2)
    doc.add_paragraph(
        "Neither Party may assign or transfer this Agreement or any of its rights "
        "or obligations hereunder without the prior written consent of the other "
        "Party, except that either Party may assign this Agreement to an affiliate "
        "or in connection with a merger, acquisition, or sale of all or substantially "
        "all of its assets, provided that the assignee assumes all obligations under "
        "this Agreement."
    )

    doc.add_heading("18.5 Notices", level=2)
    doc.add_paragraph(
        "All notices under this Agreement shall be in writing and delivered by "
        "certified mail, overnight courier, or email (with confirmation of receipt) "
        "to the addresses specified in Schedule F, Section 5. Notices shall be deemed "
        "received: (a) if by certified mail, five (5) Business Days after mailing; "
        "(b) if by overnight courier, one (1) Business Day after dispatch; or "
        "(c) if by email, upon confirmation of receipt."
    )

    # ── Schedule A: Technology Assets ──
    doc.add_heading("Schedule A: Technology Assets", level=1)

    doc.add_heading("Part I: Infrastructure Components", level=2)
    infra_table = doc.add_table(rows=11, cols=5)
    infra_table.style = "Table Grid"
    infra_headers = ["Asset Category", "Description", "Quantity", "Criticality", "Location"]
    for i, h in enumerate(infra_headers):
        infra_table.rows[0].cells[i].text = h
    infra_assets = [
        ("Physical Servers", "Dell PowerEdge R750", "480", "Critical", "All DCs"),
        ("Virtual Servers", "VMware ESXi 8.0 hosts", "1,920", "Critical", "All DCs"),
        ("Network Switches", "Cisco Nexus 9000 series", "240", "Critical", "All locations"),
        ("Firewalls", "Palo Alto PA-5400 series", "48", "Critical", "All DCs"),
        ("Load Balancers", "F5 BIG-IP i10800", "24", "High", "Primary DCs"),
        ("SAN Storage", "NetApp AFF A800", "12 arrays", "Critical", "Primary DCs"),
        ("NAS Storage", "NetApp FAS9500", "6 arrays", "High", "Primary DCs"),
        ("Backup Systems", "Veeam + Dell Data Domain", "12 units", "High", "All DCs"),
        ("Desktop/Laptop", "Dell Latitude 5540 / Inspiron 16", "15,000", "Medium", "All offices"),
        ("Printers/MFDs", "HP LaserJet Enterprise M611", "850", "Low", "All offices"),
    ]
    for row_idx, (cat, desc, qty, crit, loc) in enumerate(infra_assets, start=1):
        infra_table.rows[row_idx].cells[0].text = cat
        infra_table.rows[row_idx].cells[1].text = desc
        infra_table.rows[row_idx].cells[2].text = qty
        infra_table.rows[row_idx].cells[3].text = crit
        infra_table.rows[row_idx].cells[4].text = loc

    doc.add_heading("Part II: Enterprise Applications", level=2)
    app_table = doc.add_table(rows=9, cols=4)
    app_table.style = "Table Grid"
    app_headers = ["Application", "Version", "Users", "Support Level"]
    for i, h in enumerate(app_headers):
        app_table.rows[0].cells[i].text = h
    app_data = [
        ("SAP S/4HANA", "2023 FPS02", "4,500", "L2/L3"),
        ("Salesforce CRM", "Enterprise Edition", "2,800", "L2/L3"),
        ("Oracle E-Business Suite", "R12.2.12", "1,200", "L2/L3"),
        ("MeridianTrade (Custom)", "v4.7.2", "350", "L2/L3"),
        ("ServiceNow", "Utah release", "8,000", "L2/L3"),
        ("Microsoft 365", "E5 license", "15,000", "L1/L2"),
        ("Tableau Server", "2024.1", "600", "L2"),
        ("MuleSoft Anypoint", "4.5", "150", "L2/L3"),
    ]
    for row_idx, (app, ver, users, level) in enumerate(app_data, start=1):
        app_table.rows[row_idx].cells[0].text = app
        app_table.rows[row_idx].cells[1].text = ver
        app_table.rows[row_idx].cells[2].text = users
        app_table.rows[row_idx].cells[3].text = level

    # ── Schedule B: Service Locations ──
    doc.add_heading("Schedule B: Service Locations", level=1)
    loc_table = doc.add_table(rows=13, cols=5)
    loc_table.style = "Table Grid"
    loc_headers = ["Location", "City", "Country", "Type", "FTE Count"]
    for i, h in enumerate(loc_headers):
        loc_table.rows[0].cells[i].text = h
    locations = [
        ("NYC-HQ", "New York", "United States", "Headquarters", "45"),
        ("NYC-DC1", "New York", "United States", "Data Center (Tier IV)", "20"),
        ("CHI-OFF", "Chicago", "United States", "Regional Office", "15"),
        ("LON-OFF", "London", "United Kingdom", "Regional Office", "25"),
        ("LON-DC2", "London", "United Kingdom", "Data Center (Tier III)", "15"),
        ("FRA-OFF", "Frankfurt", "Germany", "Regional Office", "10"),
        ("SIN-OFF", "Singapore", "Singapore", "Regional Office", "20"),
        ("HYD-ODC", "Hyderabad", "India", "Offshore Dev Center", "120"),
        ("BLR-ODC", "Bangalore", "India", "Offshore Dev Center", "60"),
        ("PUN-ODC", "Pune", "India", "Offshore Dev Center", "40"),
        ("MUM-DC3", "Mumbai", "India", "Data Center (Tier III)", "10"),
        ("TOK-OFF", "Tokyo", "Japan", "Regional Office", "10"),
    ]
    for row_idx, (code, city, country, typ, fte) in enumerate(locations, start=1):
        loc_table.rows[row_idx].cells[0].text = code
        loc_table.rows[row_idx].cells[1].text = city
        loc_table.rows[row_idx].cells[2].text = country
        loc_table.rows[row_idx].cells[3].text = typ
        loc_table.rows[row_idx].cells[4].text = fte

    # ── Schedule C: SLA Penalty Matrix ──
    doc.add_heading("Schedule C: SLA Penalty Matrix", level=1)
    doc.add_paragraph(
        "The following penalty matrix applies when Service Levels are not met. "
        "Penalties are cumulative and calculated monthly."
    )
    penalty_table = doc.add_table(rows=6, cols=4)
    penalty_table.style = "Table Grid"
    pen_headers = ["Metric", "Target", "Penalty Trigger", "Penalty Amount"]
    for i, h in enumerate(pen_headers):
        penalty_table.rows[0].cells[i].text = h
    penalties = [
        ("System Availability (Critical)", "99.99%", "Below 99.95%", "2% of monthly Base Fee per 0.01% shortfall"),
        ("Incident Response (Sev 1)", "15 minutes", "Exceeds 30 minutes", "$10,000 per incident"),
        ("Incident Resolution (Sev 1)", "4 hours", "Exceeds 8 hours", "$25,000 per incident"),
        ("Change Success Rate", "98%", "Below 95%", "1% of monthly Base Fee"),
        ("Customer Satisfaction (CSAT)", "4.2/5.0", "Below 3.8/5.0", "1.5% of monthly Base Fee"),
    ]
    for row_idx, (metric, target, trigger, amount) in enumerate(penalties, start=1):
        penalty_table.rows[row_idx].cells[0].text = metric
        penalty_table.rows[row_idx].cells[1].text = target
        penalty_table.rows[row_idx].cells[2].text = trigger
        penalty_table.rows[row_idx].cells[3].text = amount

    # ── Schedule D: Compliance Requirements ──
    doc.add_heading("Schedule D: Compliance Requirements", level=1)
    compliance_table = doc.add_table(rows=7, cols=3)
    compliance_table.style = "Table Grid"
    comp_headers = ["Standard/Regulation", "Scope", "Certification Required By"]
    for i, h in enumerate(comp_headers):
        compliance_table.rows[0].cells[i].text = h
    compliance = [
        ("SOC 2 Type II", "All service delivery operations", "Within 6 months of Effective Date"),
        ("ISO 27001:2022", "Information security management", "Within 12 months of Effective Date"),
        ("PCI DSS v4.0", "Payment processing systems only", "Prior to handling payment data"),
        ("GDPR", "All EU personal data processing", "From Effective Date"),
        ("CCPA", "California consumer data", "From Effective Date"),
        ("HIPAA", "Healthcare data (if applicable)", "Prior to handling healthcare data"),
    ]
    for row_idx, (std, scope, cert_by) in enumerate(compliance, start=1):
        compliance_table.rows[row_idx].cells[0].text = std
        compliance_table.rows[row_idx].cells[1].text = scope
        compliance_table.rows[row_idx].cells[2].text = cert_by

    # Save
    doc.save(FIXTURES_DIR / "complex_it_outsourcing.docx")
    print("Created complex_it_outsourcing.docx")


def create_procurement_framework_pdf() -> None:
    """Create a complex procurement framework agreement PDF.

    Simulates a real-world multi-category procurement framework with:
    - Tiered pricing with volume discounts
    - Performance bonds and bank guarantees
    - Liquidated damages
    - Delivery schedules with milestones
    - Quality assurance requirements
    - Sustainability and ESG clauses
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, KeepTogether,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.lib.units import inch
    except ImportError:
        print("reportlab not installed — skipping PDF fixture creation")
        return

    pdf_path = FIXTURES_DIR / "complex_procurement_framework.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(ParagraphStyle(
        name='ContractTitle',
        parent=styles['Title'],
        fontSize=18,
        spaceAfter=20,
    ))
    styles.add(ParagraphStyle(
        name='SectionHead',
        parent=styles['Heading1'],
        fontSize=14,
        spaceBefore=16,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name='SubSection',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name='ContractBody',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    ))

    story = []

    # Title
    story.append(Paragraph("PROCUREMENT FRAMEWORK AGREEMENT", styles['ContractTitle']))
    story.append(Paragraph("Contract No: PFA-2025-1203", styles['ContractBody']))
    story.append(Paragraph("Classification: CONFIDENTIAL — COMMERCIAL IN CONFIDENCE", styles['ContractBody']))
    story.append(Spacer(1, 20))

    # Parties
    story.append(Paragraph("1. PARTIES AND RECITALS", styles['SectionHead']))
    story.append(Paragraph(
        '1.1 This Procurement Framework Agreement ("Framework Agreement" or "PFA") '
        'is entered into as of April 15, 2025 (the "Effective Date") by and between:',
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        '<b>Pinnacle Manufacturing Group plc</b>, a public limited company incorporated '
        'in England and Wales (Company No. 04829173), with its registered office at '
        '100 Thames Street, London EC4R 1DD, United Kingdom (hereinafter referred to '
        'as the "Buyer" or "Pinnacle");',
        styles['ContractBody'],
    ))
    story.append(Paragraph("and", styles['ContractBody']))
    story.append(Paragraph(
        '<b>GlobalSource Industrial Supply Co., Ltd.</b>, a company incorporated under '
        'the laws of the Republic of Singapore (UEN: 201534892K), with its registered '
        'office at 1 Raffles Place, #44-02 One Raffles Place Tower 2, Singapore 048616 '
        '(hereinafter referred to as the "Supplier" or "GlobalSource");',
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        '(each a "Party" and collectively the "Parties").',
        styles['ContractBody'],
    ))
    story.append(Spacer(1, 8))

    story.append(Paragraph("1.2 Recitals", styles['SubSection']))
    story.append(Paragraph(
        "WHEREAS the Buyer operates manufacturing facilities across fourteen (14) "
        "countries and requires a reliable, cost-effective supply chain for industrial "
        "raw materials, components, and maintenance supplies;",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "WHEREAS the Supplier is a leading global distributor of industrial supplies "
        "with operations in twenty-three (23) countries and a catalog of over 500,000 "
        "stock-keeping units (SKUs);",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "WHEREAS the Parties wish to establish a framework for the procurement of "
        "goods and related services over a three-year period with options for extension;",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "NOW, THEREFORE, in consideration of the mutual covenants and agreements "
        "herein contained, the Parties agree as follows:",
        styles['ContractBody'],
    ))

    # Section 2: Definitions
    story.append(Paragraph("2. DEFINITIONS", styles['SectionHead']))
    defs = [
        ('"Authorized Buyer"', 'means Pinnacle and any of its subsidiaries or affiliates listed in Annex 1 that are authorized to place Purchase Orders under this Framework Agreement.'),
        ('"Catalogue"', 'means the Supplier\'s product catalogue as updated from time to time and made available through the Supplier\'s electronic procurement portal.'),
        ('"Delivery Point"', 'means the location specified in a Purchase Order where goods are to be delivered, as listed in Annex 2.'),
        ('"Goods"', 'means the products, materials, components, and supplies described in the Categories set out in Annex 3.'),
        ('"Key Performance Indicator" or "KPI"', 'means the performance metrics set out in Annex 5 against which the Supplier\'s performance shall be measured.'),
        ('"Liquidated Damages"', 'means the pre-estimated damages payable by the Supplier for failure to meet delivery or quality obligations, as specified in Section 8.'),
        ('"Performance Bond"', 'means the bank guarantee or performance bond required under Section 10.'),
        ('"Purchase Order" or "PO"', 'means a written order placed by an Authorized Buyer for Goods and/or Services under this Framework Agreement.'),
        ('"Services"', 'means any installation, commissioning, training, or maintenance services ordered in connection with the Goods.'),
        ('"Total Framework Value"', 'means the estimated aggregate value of all Purchase Orders over the Term, currently estimated at GBP 85,000,000 (Eighty-Five Million British Pounds).'),
    ]
    for term, defn in defs:
        story.append(Paragraph(f"<b>{term}</b> {defn}", styles['ContractBody']))

    # Section 3: Term
    story.append(Paragraph("3. TERM AND EXTENSION", styles['SectionHead']))
    story.append(Paragraph(
        "3.1 This Framework Agreement shall commence on the Effective Date and "
        "continue for an initial period of three (3) years (the \"Initial Term\"), "
        "expiring on April 14, 2028.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "3.2 The Buyer may, at its sole discretion, extend this Framework Agreement "
        "for up to two (2) additional periods of one (1) year each (each an "
        "\"Extension Period\"), by providing written notice to the Supplier at least "
        "ninety (90) days before the expiry of the then-current term.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "3.3 The maximum total duration of this Framework Agreement, including all "
        "extensions, shall not exceed five (5) years from the Effective Date.",
        styles['ContractBody'],
    ))

    # Section 4: Pricing
    story.append(Paragraph("4. PRICING AND VOLUME DISCOUNTS", styles['SectionHead']))
    story.append(Paragraph(
        "4.1 Pricing for Goods shall be as set out in the Catalogue, subject to the "
        "volume discounts specified in this Section 4 and the category-specific pricing "
        "in Annex 3.",
        styles['ContractBody'],
    ))
    story.append(Paragraph("4.2 Volume Discount Tiers", styles['SubSection']))
    story.append(Paragraph(
        "The following volume discounts shall apply to the aggregate annual spend "
        "across all Categories:",
        styles['ContractBody'],
    ))

    # Volume discount table
    vol_data = [
        ["Annual Spend Tier", "Discount Rate", "Rebate Frequency"],
        ["Up to GBP 5,000,000", "Catalogue price (no discount)", "N/A"],
        ["GBP 5,000,001 — GBP 15,000,000", "5% off Catalogue price", "Quarterly"],
        ["GBP 15,000,001 — GBP 25,000,000", "8% off Catalogue price", "Quarterly"],
        ["GBP 25,000,001 — GBP 40,000,000", "12% off Catalogue price", "Monthly"],
        ["Above GBP 40,000,000", "15% off Catalogue price + 2% annual rebate", "Monthly"],
    ]
    vol_table = Table(vol_data, colWidths=[2.2*inch, 2.5*inch, 1.3*inch])
    vol_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
    ]))
    story.append(vol_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph(
        "4.3 The Supplier shall not increase Catalogue prices by more than the "
        "lesser of: (a) three percent (3%) per annum; or (b) the UK Consumer Price "
        "Index (CPI) increase for the preceding twelve (12) months as published by "
        "the Office for National Statistics. Price increases shall take effect only "
        "once per year, on the anniversary of the Effective Date, and shall require "
        "sixty (60) days' prior written notice.",
        styles['ContractBody'],
    ))

    story.append(Paragraph(
        "4.4 For bespoke or non-catalogue items, the Supplier shall provide a "
        "quotation within five (5) Business Days of the Buyer's request. Quotations "
        "shall remain valid for thirty (30) days from the date of issue.",
        styles['ContractBody'],
    ))

    # Section 5: Ordering
    story.append(Paragraph("5. ORDERING AND DELIVERY", styles['SectionHead']))
    story.append(Paragraph(
        "5.1 Purchase Orders shall be placed electronically through the Supplier's "
        "procurement portal or via email to the designated order desk. Each Purchase "
        "Order shall specify: (a) the Goods required by SKU and quantity; (b) the "
        "Delivery Point; (c) the required delivery date; and (d) any special handling "
        "or packaging requirements.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "5.2 The Supplier shall acknowledge receipt of each Purchase Order within "
        "one (1) Business Day and confirm the delivery schedule within two (2) "
        "Business Days.",
        styles['ContractBody'],
    ))
    story.append(Paragraph("5.3 Delivery Lead Times", styles['SubSection']))

    lead_data = [
        ["Category", "Standard Lead Time", "Express Lead Time (surcharge)"],
        ["Category A: Raw Materials", "10 Business Days", "5 Business Days (+15%)"],
        ["Category B: Mechanical Components", "15 Business Days", "7 Business Days (+20%)"],
        ["Category C: Electrical Components", "12 Business Days", "6 Business Days (+18%)"],
        ["Category D: Safety Equipment", "8 Business Days", "3 Business Days (+25%)"],
        ["Category E: MRO Supplies", "5 Business Days", "2 Business Days (+10%)"],
        ["Category F: Packaging Materials", "7 Business Days", "3 Business Days (+12%)"],
    ]
    lead_table = Table(lead_data, colWidths=[2*inch, 2*inch, 2.5*inch])
    lead_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
    ]))
    story.append(lead_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph(
        "5.4 All deliveries shall be made DDP (Delivered Duty Paid, Incoterms 2020) "
        "to the specified Delivery Point. The Supplier shall bear all costs of "
        "transportation, insurance, customs clearance, and duties.",
        styles['ContractBody'],
    ))

    # Section 6: Quality
    story.append(Paragraph("6. QUALITY ASSURANCE", styles['SectionHead']))
    story.append(Paragraph(
        "6.1 All Goods shall conform to the specifications set out in the applicable "
        "Purchase Order and the quality standards specified in Annex 4. The Supplier "
        "shall maintain ISO 9001:2015 certification for all manufacturing facilities "
        "supplying Goods under this Framework Agreement.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "6.2 The Buyer shall have the right to inspect Goods upon delivery. The Buyer "
        "shall notify the Supplier of any defects or non-conformities within ten (10) "
        "Business Days of delivery (the \"Inspection Period\"). Failure to notify "
        "within the Inspection Period shall not constitute acceptance of defective Goods "
        "where defects are latent or not reasonably discoverable upon inspection.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "6.3 The Supplier warrants that all Goods shall be free from defects in "
        "materials and workmanship for a period of twenty-four (24) months from the "
        "date of delivery (the \"Warranty Period\"). During the Warranty Period, the "
        "Supplier shall, at its own expense, repair or replace any defective Goods "
        "within five (5) Business Days of notification.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "6.4 The Supplier shall maintain a defect rate of no more than 0.5% across "
        "all deliveries in any rolling three (3) month period. If the defect rate "
        "exceeds 1.0% in any three (3) month period, the Buyer shall have the right "
        "to: (a) require the Supplier to implement a corrective action plan within "
        "fifteen (15) Business Days; (b) suspend further Purchase Orders until the "
        "corrective action plan is approved; or (c) terminate this Framework Agreement "
        "for cause pursuant to Section 11.",
        styles['ContractBody'],
    ))

    # Section 7: Payment
    story.append(Paragraph("7. PAYMENT TERMS", styles['SectionHead']))
    story.append(Paragraph(
        "7.1 The Supplier shall issue invoices upon delivery of Goods. Each invoice "
        "shall reference the applicable Purchase Order number and shall be accompanied "
        "by proof of delivery signed by the Buyer's authorized receiving personnel.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "7.2 Payment shall be made within sixty (60) days of receipt of a valid "
        "invoice (\"Net 60\"). The Buyer may, at its discretion, opt for early "
        "payment at a discount of two percent (2%) for payment within ten (10) days "
        "(\"2/10 Net 60\").",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "7.3 The Buyer shall have the right to set off any amounts owed by the "
        "Supplier (including Liquidated Damages, warranty claims, and credit notes) "
        "against amounts payable to the Supplier.",
        styles['ContractBody'],
    ))

    # Section 8: Liquidated Damages
    story.append(Paragraph("8. LIQUIDATED DAMAGES", styles['SectionHead']))
    story.append(Paragraph(
        "8.1 The Parties acknowledge that late delivery or delivery of non-conforming "
        "Goods causes the Buyer to suffer losses that are difficult to quantify. "
        "Accordingly, the following Liquidated Damages shall apply:",
        styles['ContractBody'],
    ))

    ld_data = [
        ["Breach Type", "Liquidated Damages Rate", "Cap"],
        ["Late delivery (1-5 days)", "0.5% of PO value per day", "5% of PO value"],
        ["Late delivery (6-15 days)", "1.0% of PO value per day", "15% of PO value"],
        ["Late delivery (>15 days)", "Buyer may cancel PO + 15% of PO value", "N/A"],
        ["Quality non-conformance", "150% of replacement cost", "25% of PO value"],
        ["Documentation deficiency", "GBP 500 per missing document per day", "GBP 10,000 per PO"],
    ]
    ld_table = Table(ld_data, colWidths=[2*inch, 2.5*inch, 1.8*inch])
    ld_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#C0392B")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#FDEDEC")]),
    ]))
    story.append(ld_table)
    story.append(Spacer(1, 8))

    story.append(Paragraph(
        "8.2 The aggregate Liquidated Damages payable by the Supplier in any "
        "twelve (12) month period shall not exceed ten percent (10%) of the annual "
        "spend under this Framework Agreement. If Liquidated Damages reach this cap, "
        "the Buyer shall have the right to terminate for cause under Section 11.",
        styles['ContractBody'],
    ))

    # Section 9: Confidentiality
    story.append(Paragraph("9. CONFIDENTIALITY", styles['SectionHead']))
    story.append(Paragraph(
        "9.1 Each Party shall treat as confidential all information received from "
        "the other Party that is marked as confidential or that a reasonable person "
        "would consider confidential, including pricing, technical specifications, "
        "business plans, and customer information (\"Confidential Information\").",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "9.2 The obligations of confidentiality shall not apply to information that: "
        "(a) is or becomes publicly available through no fault of the Receiving Party; "
        "(b) was known to the Receiving Party prior to disclosure; (c) is independently "
        "developed by the Receiving Party; or (d) is required to be disclosed by law, "
        "regulation, or court order, provided that the Receiving Party gives the "
        "Disclosing Party prompt notice and cooperates in seeking protective measures.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "9.3 The confidentiality obligations under this Section shall survive "
        "termination or expiry of this Framework Agreement for a period of seven "
        "(7) years.",
        styles['ContractBody'],
    ))

    # Section 10: Performance Bond
    story.append(Paragraph("10. PERFORMANCE BOND AND GUARANTEES", styles['SectionHead']))
    story.append(Paragraph(
        "10.1 Within thirty (30) days of the Effective Date, the Supplier shall "
        "provide the Buyer with an irrevocable, unconditional bank guarantee (the "
        "\"Performance Bond\") in the amount of Five Million British Pounds "
        "(GBP 5,000,000), issued by a bank acceptable to the Buyer with a minimum "
        "credit rating of A- (S&P) or A3 (Moody's).",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "10.2 The Performance Bond shall remain in effect throughout the Term and "
        "for a period of six (6) months following expiry or termination. The Buyer "
        "may draw upon the Performance Bond in the event of: (a) the Supplier's "
        "failure to pay Liquidated Damages when due; (b) the Supplier's insolvency; "
        "or (c) the Supplier's material breach that remains uncured for thirty (30) "
        "days after written notice.",
        styles['ContractBody'],
    ))

    # Section 11: Termination
    story.append(Paragraph("11. TERMINATION", styles['SectionHead']))
    story.append(Paragraph(
        "11.1 <b>Termination for Convenience.</b> The Buyer may terminate this "
        "Framework Agreement at any time by providing ninety (90) days' written "
        "notice to the Supplier. Upon termination for convenience, the Buyer shall "
        "honor all Purchase Orders placed prior to the termination notice date and "
        "shall pay for all Goods delivered in accordance with this Agreement.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "11.2 <b>Termination for Cause.</b> Either Party may terminate this Framework "
        "Agreement immediately upon written notice if the other Party: (a) commits a "
        "material breach that remains uncured for thirty (30) days after written "
        "notice; (b) becomes insolvent, enters administration, or has a receiver "
        "appointed; (c) fails to maintain the Performance Bond as required under "
        "Section 10; or (d) is convicted of fraud, bribery, or corruption.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "11.3 <b>Consequences of Termination.</b> Upon termination: (a) all outstanding "
        "Purchase Orders shall be fulfilled unless otherwise agreed; (b) the Supplier "
        "shall return all Confidential Information within thirty (30) days; (c) the "
        "Supplier shall cooperate in transitioning supply to an alternative supplier "
        "for a period of up to six (6) months; and (d) all accrued rights and "
        "obligations shall survive termination.",
        styles['ContractBody'],
    ))

    # Section 12: Indemnification
    story.append(Paragraph("12. INDEMNIFICATION AND LIABILITY", styles['SectionHead']))
    story.append(Paragraph(
        "12.1 The Supplier shall indemnify and hold harmless the Buyer against all "
        "claims, damages, losses, and expenses arising from: (a) defective Goods; "
        "(b) infringement of third-party intellectual property rights; (c) breach "
        "of applicable laws or regulations; (d) personal injury or property damage "
        "caused by the Goods; or (e) the Supplier's breach of this Agreement.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "12.2 The Supplier's aggregate liability under this Framework Agreement "
        "shall not exceed one hundred and fifty percent (150%) of the Total Framework "
        "Value. This limitation shall not apply to: (a) death or personal injury "
        "caused by negligence; (b) fraud or fraudulent misrepresentation; (c) breach "
        "of confidentiality; or (d) the Supplier's indemnification obligations under "
        "Section 12.1(b) and (d).",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "12.3 Neither Party shall be liable for indirect, consequential, or special "
        "damages, except in cases of willful misconduct, fraud, or breach of "
        "confidentiality obligations.",
        styles['ContractBody'],
    ))

    # Section 13: Sustainability
    story.append(Paragraph("13. SUSTAINABILITY AND ESG REQUIREMENTS", styles['SectionHead']))
    story.append(Paragraph(
        "13.1 The Supplier shall comply with the Buyer's Supplier Code of Conduct "
        "(attached as Annex 6) and shall demonstrate commitment to environmental, "
        "social, and governance (\"ESG\") principles, including:",
        styles['ContractBody'],
    ))
    esg_items = [
        "Carbon footprint reduction: achieve a 30% reduction in Scope 1 and 2 emissions by 2028 (baseline: 2024)",
        "Sustainable sourcing: ensure at least 50% of raw materials are from certified sustainable sources by 2027",
        "Modern slavery: maintain zero tolerance for forced labor, child labor, and human trafficking throughout the supply chain",
        "Diversity and inclusion: report annually on workforce diversity metrics",
        "Circular economy: implement take-back or recycling programs for applicable product categories",
        "Conflict minerals: comply with the EU Conflict Minerals Regulation and Dodd-Frank Act Section 1502",
    ]
    for item in esg_items:
        story.append(Paragraph(f"• {item}", styles['ContractBody']))

    story.append(Paragraph(
        "13.2 The Supplier shall provide an annual ESG compliance report and shall "
        "permit the Buyer (or its designated auditor) to conduct sustainability "
        "audits of the Supplier's facilities and supply chain with thirty (30) days' "
        "prior notice.",
        styles['ContractBody'],
    ))

    # Section 14: Force Majeure
    story.append(Paragraph("14. FORCE MAJEURE", styles['SectionHead']))
    story.append(Paragraph(
        "14.1 Neither Party shall be liable for failure to perform its obligations "
        "under this Framework Agreement to the extent that performance is prevented "
        "by a Force Majeure Event, meaning any event beyond the reasonable control "
        "of the affected Party, including natural disasters, war, terrorism, "
        "pandemics, government sanctions, and failure of third-party infrastructure.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "14.2 The affected Party shall notify the other Party within seventy-two "
        "(72) hours and shall use all reasonable endeavors to mitigate the impact. "
        "If a Force Majeure Event continues for more than sixty (60) days, either "
        "Party may terminate this Framework Agreement upon fourteen (14) days' "
        "written notice.",
        styles['ContractBody'],
    ))

    # Section 15: Governing Law
    story.append(Paragraph("15. GOVERNING LAW AND JURISDICTION", styles['SectionHead']))
    story.append(Paragraph(
        "15.1 This Framework Agreement shall be governed by and construed in "
        "accordance with the laws of England and Wales.",
        styles['ContractBody'],
    ))
    story.append(Paragraph(
        "15.2 Any dispute arising out of or in connection with this Framework "
        "Agreement shall be referred to and finally resolved by arbitration under "
        "the LCIA Rules. The seat of arbitration shall be London. The tribunal "
        "shall consist of three (3) arbitrators. The language of the arbitration "
        "shall be English.",
        styles['ContractBody'],
    ))

    # Annex 2: Delivery Points
    story.append(PageBreak())
    story.append(Paragraph("ANNEX 2: DELIVERY POINTS", styles['SectionHead']))

    dp_data = [
        ["Code", "Facility", "City", "Country", "Receiving Hours"],
        ["UK-BIR", "Birmingham Manufacturing Plant", "Birmingham", "United Kingdom", "Mon-Fri 06:00-22:00"],
        ["UK-MAN", "Manchester Assembly Works", "Manchester", "United Kingdom", "Mon-Sat 07:00-19:00"],
        ["DE-STU", "Stuttgart Precision Engineering", "Stuttgart", "Germany", "Mon-Fri 06:00-18:00"],
        ["FR-LYO", "Lyon Components Facility", "Lyon", "France", "Mon-Fri 07:00-17:00"],
        ["US-DET", "Detroit Automotive Division", "Detroit", "United States", "Mon-Fri 06:00-22:00"],
        ["US-HOU", "Houston Energy Division", "Houston", "United States", "24/7"],
        ["CN-SHZ", "Shenzhen Electronics Plant", "Shenzhen", "China", "Mon-Sat 08:00-20:00"],
        ["IN-CHN", "Chennai Manufacturing Hub", "Chennai", "India", "Mon-Sat 08:00-20:00"],
        ["IN-PUN", "Pune Automotive Components", "Pune", "India", "Mon-Fri 09:00-18:00"],
        ["JP-OSA", "Osaka Quality Center", "Osaka", "Japan", "Mon-Fri 08:00-17:00"],
        ["BR-SAO", "São Paulo Distribution Center", "São Paulo", "Brazil", "Mon-Fri 07:00-19:00"],
        ["MX-MTY", "Monterrey Assembly Plant", "Monterrey", "Mexico", "Mon-Sat 06:00-22:00"],
        ["PL-WRO", "Wrocław Components Factory", "Wrocław", "Poland", "Mon-Fri 06:00-18:00"],
        ["AU-MEL", "Melbourne Warehouse", "Melbourne", "Australia", "Mon-Fri 07:00-17:00"],
    ]
    dp_table = Table(dp_data, colWidths=[0.6*inch, 2*inch, 1*inch, 1.2*inch, 1.5*inch])
    dp_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
    ]))
    story.append(dp_table)

    # Annex 3: Product Categories
    story.append(Spacer(1, 16))
    story.append(Paragraph("ANNEX 3: PRODUCT CATEGORIES AND PRICING", styles['SectionHead']))

    cat_data = [
        ["Category", "Description", "Est. Annual Spend", "Discount Tier"],
        ["A: Raw Materials", "Steel, aluminum, copper, polymers, composites", "GBP 12,000,000", "Tier 3 (8%)"],
        ["B: Mechanical Components", "Bearings, gears, fasteners, seals, hydraulics", "GBP 8,500,000", "Tier 2 (5%)"],
        ["C: Electrical Components", "Motors, drives, sensors, PLCs, cabling", "GBP 6,200,000", "Tier 2 (5%)"],
        ["D: Safety Equipment", "PPE, fire suppression, gas detection, barriers", "GBP 2,800,000", "Tier 1 (0%)"],
        ["E: MRO Supplies", "Lubricants, cleaning, tools, consumables", "GBP 4,100,000", "Tier 1 (0%)"],
        ["F: Packaging Materials", "Pallets, wrapping, containers, labels", "GBP 1,900,000", "Tier 1 (0%)"],
    ]
    cat_table = Table(cat_data, colWidths=[1.5*inch, 2.2*inch, 1.3*inch, 1.2*inch])
    cat_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2C3E50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F8F9FA")]),
    ]))
    story.append(cat_table)

    # Annex 5: KPIs
    story.append(Spacer(1, 16))
    story.append(Paragraph("ANNEX 5: KEY PERFORMANCE INDICATORS", styles['SectionHead']))

    kpi_data = [
        ["KPI", "Target", "Measurement", "Consequence"],
        ["On-Time Delivery (OTD)", "≥ 95%", "Monthly", "LD per Section 8 if < 90%"],
        ["Order Accuracy", "≥ 99%", "Monthly", "Corrective action if < 97%"],
        ["Defect Rate", "≤ 0.5%", "Quarterly (rolling)", "Suspension if > 1.0%"],
        ["Invoice Accuracy", "≥ 98%", "Monthly", "Payment withheld until corrected"],
        ["Emergency Response Time", "< 4 hours", "Per incident", "LD if > 8 hours"],
        ["Sustainability Score", "≥ 80/100", "Annual audit", "Improvement plan required if < 70"],
        ["Customer Satisfaction (NPS)", "≥ 50", "Semi-annual survey", "Review meeting if < 40"],
    ]
    kpi_table = Table(kpi_data, colWidths=[1.5*inch, 1*inch, 1.2*inch, 2.5*inch])
    kpi_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#27AE60")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EAFAF1")]),
    ]))
    story.append(kpi_table)

    doc.build(story)
    print("Created complex_procurement_framework.pdf")


if __name__ == "__main__":
    create_it_outsourcing_agreement_docx()
    create_procurement_framework_pdf()
