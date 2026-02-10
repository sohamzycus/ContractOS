"""Generate LegalBench / CUAD-style contract fixtures for benchmark testing.

These documents simulate real-world contracts covering the key categories
evaluated by the LegalBench benchmark (https://www.legalevalhub.ai/):

1. **LegalBench NDA** — Bilateral Non-Disclosure Agreement
   Covers: contract_nli_confidentiality, contract_nli_sharing_with_employees,
   contract_nli_survival_of_obligations, definition_classification,
   definition_extraction

2. **CUAD License Agreement** — Software License & Distribution Agreement
   Covers: cuad_license_grant, cuad_non-compete, cuad_termination_for_convenience,
   cuad_cap_on_liability, cuad_governing_law, cuad_insurance,
   cuad_ip_ownership_assignment, cuad_audit_rights, cuad_renewal_term

Usage:
    python tests/fixtures/create_legalbench_fixtures.py
"""

from __future__ import annotations

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Install python-docx: pip install python-docx")
    sys.exit(1)


FIXTURES_DIR = Path(__file__).parent


# ─────────────────────────────────────────────────────────────────────
# 1. LegalBench-style Bilateral NDA
# ─────────────────────────────────────────────────────────────────────

def create_legalbench_nda() -> Path:
    """Create a bilateral NDA covering LegalBench contract_nli categories."""
    doc = Document()

    # Title
    title = doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement (the \"Agreement\") is entered into "
        "as of March 15, 2025 (the \"Effective Date\") by and between:"
    )

    # Parties
    doc.add_paragraph(
        "1. Nexus Dynamics Inc., a Delaware corporation with its principal offices "
        "at 2400 Innovation Drive, Suite 800, San Jose, California 95134 "
        "(hereinafter referred to as \"Nexus\" or \"Disclosing Party\"); and"
    )
    doc.add_paragraph(
        "2. Quantum Leap Technologies Ltd., a company incorporated under the laws "
        "of England and Wales with registered office at 15 Canary Wharf, Tower 3, "
        "London E14 5AB, United Kingdom "
        "(hereinafter referred to as \"Quantum\" or \"Receiving Party\")."
    )
    doc.add_paragraph(
        "Nexus and Quantum are each referred to individually as a \"Party\" and "
        "collectively as the \"Parties\"."
    )

    # RECITALS
    doc.add_heading("RECITALS", level=1)
    doc.add_paragraph(
        "WHEREAS, the Parties wish to explore a potential business relationship "
        "concerning the development of advanced quantum computing middleware "
        "(the \"Purpose\"); and"
    )
    doc.add_paragraph(
        "WHEREAS, in connection with the Purpose, each Party may disclose to the "
        "other certain confidential and proprietary information; and"
    )
    doc.add_paragraph(
        "WHEREAS, the Parties desire to establish the terms and conditions under "
        "which such information will be disclosed and protected."
    )
    doc.add_paragraph(
        "NOW, THEREFORE, in consideration of the mutual covenants and agreements "
        "contained herein, and for other good and valuable consideration, the "
        "receipt and sufficiency of which are hereby acknowledged, the Parties "
        "agree as follows:"
    )

    # 1. Definitions
    doc.add_heading("1. Definitions", level=1)
    doc.add_paragraph(
        "1.1 \"Confidential Information\" means any and all non-public, proprietary, "
        "or confidential information disclosed by either Party to the other Party, "
        "whether orally, in writing, electronically, or by inspection of tangible "
        "objects, including but not limited to: (a) trade secrets, inventions, ideas, "
        "processes, formulas, source and object code, data, programs, software, "
        "other works of authorship, know-how, improvements, discoveries, developments, "
        "designs, and techniques; (b) information regarding plans for research, "
        "development, new products, marketing and selling, business plans, budgets "
        "and unpublished financial statements, licenses, prices and costs, suppliers, "
        "and customers; and (c) information regarding the skills and compensation of "
        "employees, contractors, and agents of the Disclosing Party."
    )
    doc.add_paragraph(
        "1.2 \"Disclosing Party\" means the Party disclosing Confidential Information."
    )
    doc.add_paragraph(
        "1.3 \"Receiving Party\" means the Party receiving Confidential Information."
    )
    doc.add_paragraph(
        "1.4 \"Representatives\" means the officers, directors, employees, agents, "
        "advisors (including attorneys, accountants, consultants, bankers, and "
        "financial advisors), and affiliates of a Party."
    )
    doc.add_paragraph(
        "1.5 \"Permitted Purpose\" means the evaluation, negotiation, and "
        "implementation of the potential business relationship described in the "
        "Recitals."
    )

    # 2. Confidentiality Obligations
    doc.add_heading("2. Confidentiality Obligations", level=1)
    doc.add_paragraph(
        "2.1 The Receiving Party shall: (a) hold the Confidential Information in "
        "strict confidence; (b) not disclose the Confidential Information to any "
        "third party without the prior written consent of the Disclosing Party; "
        "(c) use the Confidential Information solely for the Permitted Purpose; "
        "and (d) protect the Confidential Information using the same degree of care "
        "it uses to protect its own confidential information, but in no event less "
        "than reasonable care."
    )
    doc.add_paragraph(
        "2.2 The Receiving Party may disclose Confidential Information to its "
        "Representatives who (a) have a need to know such information for the "
        "Permitted Purpose, (b) have been informed of the confidential nature of "
        "such information, and (c) are bound by written confidentiality obligations "
        "no less restrictive than those contained herein. The Receiving Party shall "
        "be responsible for any breach of this Agreement by its Representatives."
    )
    doc.add_paragraph(
        "2.3 The Receiving Party shall not reverse engineer, disassemble, or "
        "decompile any prototypes, software, samples, or other tangible objects "
        "that embody the Disclosing Party's Confidential Information."
    )

    # 3. Exclusions
    doc.add_heading("3. Exclusions from Confidential Information", level=1)
    doc.add_paragraph(
        "3.1 Confidential Information shall not include information that: "
        "(a) was publicly known and made generally available in the public domain "
        "prior to the time of disclosure by the Disclosing Party; "
        "(b) becomes publicly known and made generally available after disclosure "
        "by the Disclosing Party through no wrongful action or inaction of the "
        "Receiving Party; "
        "(c) is already in the possession of the Receiving Party at the time of "
        "disclosure, as shown by the Receiving Party's files and records; "
        "(d) is obtained by the Receiving Party from a third party without a breach "
        "of such third party's obligations of confidentiality; or "
        "(e) is independently developed by the Receiving Party without use of or "
        "reference to the Disclosing Party's Confidential Information."
    )

    # 4. Compelled Disclosure
    doc.add_heading("4. Compelled Disclosure", level=1)
    doc.add_paragraph(
        "4.1 If the Receiving Party is compelled by law, regulation, or legal "
        "process to disclose any Confidential Information, the Receiving Party "
        "shall provide the Disclosing Party with prompt written notice of such "
        "requirement so that the Disclosing Party may seek a protective order or "
        "other appropriate remedy. The Receiving Party shall disclose only that "
        "portion of the Confidential Information that it is legally required to "
        "disclose and shall use commercially reasonable efforts to obtain "
        "confidential treatment for any Confidential Information so disclosed."
    )

    # 5. Return of Materials
    doc.add_heading("5. Return of Materials", level=1)
    doc.add_paragraph(
        "5.1 Upon the written request of the Disclosing Party or upon termination "
        "of this Agreement, the Receiving Party shall promptly return or destroy "
        "all documents, materials, and other tangible manifestations of Confidential "
        "Information in its possession or in the possession of its Representatives, "
        "and shall provide written certification of such return or destruction. "
        "Notwithstanding the foregoing, the Receiving Party may retain one (1) "
        "archival copy of the Confidential Information solely for legal compliance "
        "and audit purposes, subject to the continuing obligations of this Agreement."
    )

    # 6. No License
    doc.add_heading("6. No License or Warranty", level=1)
    doc.add_paragraph(
        "6.1 Nothing in this Agreement grants the Receiving Party any license or "
        "right to use the Confidential Information except as expressly set forth "
        "herein. All Confidential Information remains the property of the "
        "Disclosing Party. No license or conveyance of any intellectual property "
        "rights is granted or implied by this Agreement."
    )
    doc.add_paragraph(
        "6.2 ALL CONFIDENTIAL INFORMATION IS PROVIDED \"AS IS\" WITHOUT WARRANTY "
        "OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO WARRANTIES "
        "OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE, OR NON-INFRINGEMENT."
    )

    # 7. Term and Termination
    doc.add_heading("7. Term and Termination", level=1)
    doc.add_paragraph(
        "7.1 This Agreement shall commence on the Effective Date and shall remain "
        "in effect for a period of three (3) years, unless earlier terminated by "
        "either Party upon thirty (30) days' prior written notice to the other Party."
    )
    doc.add_paragraph(
        "7.2 The confidentiality obligations set forth in Section 2 shall survive "
        "termination or expiration of this Agreement for a period of five (5) years "
        "from the date of disclosure of the applicable Confidential Information."
    )

    # 8. Remedies
    doc.add_heading("8. Remedies", level=1)
    doc.add_paragraph(
        "8.1 The Parties acknowledge that any breach of this Agreement may cause "
        "irreparable harm to the Disclosing Party for which monetary damages would "
        "be an inadequate remedy. Accordingly, the Disclosing Party shall be "
        "entitled to seek equitable relief, including injunction and specific "
        "performance, in addition to all other remedies available at law or in "
        "equity, without the necessity of proving actual damages or posting a bond."
    )
    doc.add_paragraph(
        "8.2 In the event of a breach, the breaching Party shall indemnify and "
        "hold harmless the non-breaching Party from and against all losses, damages, "
        "liabilities, costs, and expenses (including reasonable attorneys' fees) "
        "arising out of or relating to such breach. The total aggregate liability "
        "of either Party under this Agreement shall not exceed Five Million United "
        "States Dollars (USD $5,000,000)."
    )

    # 9. Non-Solicitation
    doc.add_heading("9. Non-Solicitation", level=1)
    doc.add_paragraph(
        "9.1 During the term of this Agreement and for a period of twelve (12) "
        "months following its termination, neither Party shall directly or "
        "indirectly solicit, recruit, or hire any employee or contractor of the "
        "other Party who was involved in the Permitted Purpose, without the prior "
        "written consent of the other Party."
    )

    # 10. Governing Law
    doc.add_heading("10. Governing Law and Dispute Resolution", level=1)
    doc.add_paragraph(
        "10.1 This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of Delaware, United States of America, without "
        "regard to its conflict of laws principles."
    )
    doc.add_paragraph(
        "10.2 Any dispute arising out of or in connection with this Agreement shall "
        "be resolved by binding arbitration administered by the American Arbitration "
        "Association (\"AAA\") in accordance with its Commercial Arbitration Rules. "
        "The arbitration shall be conducted in New York, New York, by a panel of "
        "three (3) arbitrators. The language of the arbitration shall be English."
    )

    # 11. General Provisions
    doc.add_heading("11. General Provisions", level=1)
    doc.add_paragraph(
        "11.1 Entire Agreement. This Agreement constitutes the entire agreement "
        "between the Parties with respect to the subject matter hereof and "
        "supersedes all prior and contemporaneous agreements, understandings, "
        "negotiations, and discussions, whether oral or written."
    )
    doc.add_paragraph(
        "11.2 Amendment. This Agreement may not be amended or modified except by "
        "a written instrument signed by both Parties."
    )
    doc.add_paragraph(
        "11.3 Assignment. Neither Party may assign or transfer this Agreement or "
        "any rights or obligations hereunder without the prior written consent of "
        "the other Party, except in connection with a merger, acquisition, or sale "
        "of all or substantially all of its assets."
    )
    doc.add_paragraph(
        "11.4 Severability. If any provision of this Agreement is held to be "
        "invalid or unenforceable, the remaining provisions shall continue in "
        "full force and effect."
    )
    doc.add_paragraph(
        "11.5 Waiver. The failure of either Party to enforce any provision of this "
        "Agreement shall not constitute a waiver of such provision or the right to "
        "enforce it at a later time."
    )
    doc.add_paragraph(
        "11.6 Notices. All notices under this Agreement shall be in writing and "
        "shall be deemed given when delivered personally, sent by confirmed email, "
        "or sent by certified mail, return receipt requested, to the addresses set "
        "forth above."
    )
    doc.add_paragraph(
        "11.7 Counterparts. This Agreement may be executed in counterparts, each "
        "of which shall be deemed an original, and all of which together shall "
        "constitute one and the same instrument."
    )

    # Signature block
    doc.add_paragraph("")
    doc.add_paragraph("IN WITNESS WHEREOF, the Parties have executed this Agreement "
                       "as of the Effective Date.")
    doc.add_paragraph("")

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "NEXUS DYNAMICS INC."
    table.cell(0, 1).text = "QUANTUM LEAP TECHNOLOGIES LTD."
    table.cell(1, 0).text = "By: ________________________"
    table.cell(1, 1).text = "By: ________________________"
    table.cell(2, 0).text = "Name: Dr. Sarah Chen, CEO"
    table.cell(2, 1).text = "Name: James Harrington, Managing Director"
    table.cell(3, 0).text = "Date: March 15, 2025"
    table.cell(3, 1).text = "Date: March 15, 2025"

    out = FIXTURES_DIR / "legalbench_nda.docx"
    doc.save(str(out))
    print(f"Created: {out}  ({out.stat().st_size:,} bytes)")
    return out


# ─────────────────────────────────────────────────────────────────────
# 2. CUAD-style Software License & Distribution Agreement
# ─────────────────────────────────────────────────────────────────────

def create_cuad_license_agreement() -> Path:
    """Create a CUAD-style license agreement covering key CUAD categories."""
    doc = Document()

    title = doc.add_heading("SOFTWARE LICENSE AND DISTRIBUTION AGREEMENT", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Agreement No.: CUAD-SLA-2025-0042\n"
        "Effective Date: January 1, 2025"
    )

    doc.add_paragraph(
        "This Software License and Distribution Agreement (the \"Agreement\") is "
        "entered into by and between:"
    )
    doc.add_paragraph(
        "Licensor: Apex Software Corporation, a California corporation, with its "
        "principal place of business at 1 Infinite Loop, Building 7, Cupertino, "
        "California 95014 (\"Apex\" or \"Licensor\"); and"
    )
    doc.add_paragraph(
        "Licensee: GlobalTech Solutions GmbH, a company organized under the laws "
        "of Germany, with registered offices at Friedrichstraße 123, 10117 Berlin, "
        "Germany (\"GlobalTech\" or \"Licensee\")."
    )
    doc.add_paragraph(
        "Apex and GlobalTech are each referred to as a \"Party\" and collectively "
        "as the \"Parties\"."
    )

    # 1. Definitions
    doc.add_heading("1. Definitions", level=1)
    definitions = [
        ("\"Licensed Software\"", "means the Apex Enterprise Platform version 12.x, "
         "including all modules, components, documentation, updates, and patches "
         "provided by Licensor during the term of this Agreement."),
        ("\"Territory\"", "means the European Economic Area (EEA), the United Kingdom, "
         "and Switzerland."),
        ("\"End Users\"", "means the customers and clients of Licensee who are "
         "authorized to use the Licensed Software pursuant to sublicenses granted "
         "by Licensee in accordance with this Agreement."),
        ("\"Intellectual Property Rights\"", "means all patents, copyrights, trademarks, "
         "trade secrets, and other intellectual property rights recognized in any "
         "jurisdiction worldwide."),
        ("\"Derivative Works\"", "means any modification, enhancement, adaptation, "
         "translation, or other work based upon the Licensed Software."),
        ("\"Annual License Fee\"", "means the fee of EUR 2,400,000 (Two Million Four "
         "Hundred Thousand Euros) per annum, payable in quarterly installments of "
         "EUR 600,000."),
        ("\"Minimum Commitment\"", "means the minimum annual revenue commitment of "
         "EUR 1,200,000 in sublicense fees that Licensee must generate from End Users "
         "during each contract year."),
    ]
    for term, defn in definitions:
        doc.add_paragraph(f"{term} {defn}")

    # 2. License Grant
    doc.add_heading("2. License Grant", level=1)
    doc.add_paragraph(
        "2.1 Grant of License. Subject to the terms and conditions of this Agreement "
        "and payment of the Annual License Fee, Licensor hereby grants to Licensee "
        "a non-exclusive, non-transferable, revocable license to: "
        "(a) use, copy, and install the Licensed Software on Licensee's servers "
        "within the Territory; "
        "(b) distribute and sublicense the Licensed Software to End Users within "
        "the Territory; "
        "(c) provide support and maintenance services to End Users in connection "
        "with the Licensed Software; and "
        "(d) create Derivative Works for the sole purpose of integrating the "
        "Licensed Software with Licensee's proprietary systems."
    )
    doc.add_paragraph(
        "2.2 Restrictions. Licensee shall not: "
        "(a) sublicense, sell, distribute, or otherwise make the Licensed Software "
        "available outside the Territory without Licensor's prior written consent; "
        "(b) reverse engineer, decompile, or disassemble the Licensed Software, "
        "except to the extent expressly permitted by applicable law; "
        "(c) remove or alter any proprietary notices, labels, or marks on the "
        "Licensed Software; or "
        "(d) use the Licensed Software for any purpose other than as expressly "
        "authorized under this Agreement."
    )
    doc.add_paragraph(
        "2.3 Affiliate License. Licensee may extend the license granted herein to "
        "its wholly-owned subsidiaries within the Territory, provided that such "
        "subsidiaries agree in writing to be bound by the terms of this Agreement. "
        "Licensee shall be jointly and severally liable for any breach by its "
        "affiliates."
    )

    # 3. Intellectual Property
    doc.add_heading("3. Intellectual Property Ownership", level=1)
    doc.add_paragraph(
        "3.1 Licensor Ownership. All right, title, and interest in and to the "
        "Licensed Software, including all Intellectual Property Rights therein, "
        "shall remain with Licensor. This Agreement does not convey to Licensee "
        "any ownership interest in the Licensed Software."
    )
    doc.add_paragraph(
        "3.2 Derivative Works. All Derivative Works created by Licensee shall be "
        "jointly owned by Licensor and Licensee. Licensee hereby assigns to "
        "Licensor an undivided fifty percent (50%) interest in all Derivative Works. "
        "Licensor shall have the right to use, license, and distribute such "
        "Derivative Works without restriction."
    )
    doc.add_paragraph(
        "3.3 Feedback. Any suggestions, enhancement requests, recommendations, or "
        "other feedback provided by Licensee regarding the Licensed Software "
        "(\"Feedback\") shall be the sole property of Licensor. Licensee hereby "
        "assigns to Licensor all right, title, and interest in and to such Feedback."
    )

    # 4. Fees and Payment
    doc.add_heading("4. Fees and Payment", level=1)
    doc.add_paragraph(
        "4.1 Annual License Fee. Licensee shall pay the Annual License Fee of "
        "EUR 2,400,000 in quarterly installments of EUR 600,000, due on the first "
        "business day of each calendar quarter."
    )
    doc.add_paragraph(
        "4.2 Revenue Sharing. In addition to the Annual License Fee, Licensee shall "
        "pay Licensor a royalty of fifteen percent (15%) of all gross revenue "
        "received by Licensee from sublicensing the Licensed Software to End Users, "
        "net of applicable taxes. Royalty payments shall be made quarterly within "
        "thirty (30) days after the end of each calendar quarter."
    )
    doc.add_paragraph(
        "4.3 Minimum Commitment. Licensee guarantees a Minimum Commitment of "
        "EUR 1,200,000 in annual sublicense revenue. If Licensee fails to meet the "
        "Minimum Commitment in any contract year, Licensee shall pay Licensor the "
        "difference between the actual sublicense revenue and the Minimum Commitment "
        "multiplied by the royalty rate (15%)."
    )
    doc.add_paragraph(
        "4.4 Price Adjustment. The Annual License Fee shall be subject to an annual "
        "increase of three percent (3%) commencing on the second anniversary of the "
        "Effective Date, unless otherwise agreed in writing by the Parties."
    )
    doc.add_paragraph(
        "4.5 Late Payment. Any amounts not paid when due shall bear interest at the "
        "rate of one and one-half percent (1.5%) per month, or the maximum rate "
        "permitted by applicable law, whichever is less."
    )

    # 5. Audit Rights
    doc.add_heading("5. Audit Rights", level=1)
    doc.add_paragraph(
        "5.1 Licensor shall have the right, upon thirty (30) days' prior written "
        "notice, to audit Licensee's books, records, and systems to verify "
        "compliance with this Agreement, including but not limited to: "
        "(a) the number of installations and End Users; "
        "(b) the accuracy of royalty payments; and "
        "(c) compliance with the license restrictions set forth in Section 2.2."
    )
    doc.add_paragraph(
        "5.2 Audits shall be conducted during normal business hours, no more than "
        "once per calendar year, by an independent certified public accounting firm "
        "selected by Licensor and reasonably acceptable to Licensee. Licensor shall "
        "bear the cost of the audit unless the audit reveals an underpayment of "
        "more than five percent (5%), in which case Licensee shall bear all audit costs."
    )

    # 6. Non-Compete
    doc.add_heading("6. Non-Competition", level=1)
    doc.add_paragraph(
        "6.1 During the term of this Agreement and for a period of twenty-four (24) "
        "months following its termination or expiration, Licensee shall not, directly "
        "or indirectly, develop, market, distribute, or license any software product "
        "that is substantially similar to or competitive with the Licensed Software "
        "within the Territory."
    )
    doc.add_paragraph(
        "6.2 The non-competition restriction in Section 6.1 shall not apply to: "
        "(a) software products that Licensee was developing or distributing prior "
        "to the Effective Date, as documented in Exhibit C; or "
        "(b) software products acquired by Licensee through a merger or acquisition, "
        "provided that Licensee divests such products within twelve (12) months of "
        "the acquisition."
    )

    # 7. Term and Termination
    doc.add_heading("7. Term and Termination", level=1)
    doc.add_paragraph(
        "7.1 Initial Term. This Agreement shall commence on the Effective Date and "
        "shall continue for an initial term of five (5) years (the \"Initial Term\")."
    )
    doc.add_paragraph(
        "7.2 Renewal. Upon expiration of the Initial Term, this Agreement shall "
        "automatically renew for successive two (2) year periods (each a \"Renewal "
        "Term\"), unless either Party provides written notice of non-renewal at "
        "least one hundred eighty (180) days prior to the expiration of the then-"
        "current term."
    )
    doc.add_paragraph(
        "7.3 Termination for Convenience. Either Party may terminate this Agreement "
        "for convenience upon ninety (90) days' prior written notice to the other "
        "Party, subject to payment of all outstanding fees and an early termination "
        "fee equal to six (6) months of the Annual License Fee."
    )
    doc.add_paragraph(
        "7.4 Termination for Cause. Either Party may terminate this Agreement "
        "immediately upon written notice if: "
        "(a) the other Party commits a material breach that remains uncured for "
        "thirty (30) days after written notice; "
        "(b) the other Party becomes insolvent, files for bankruptcy, or has a "
        "receiver appointed for its assets; or "
        "(c) the other Party is acquired by a direct competitor of the terminating "
        "Party."
    )
    doc.add_paragraph(
        "7.5 Effect of Termination. Upon termination: "
        "(a) all licenses granted herein shall immediately terminate; "
        "(b) Licensee shall cease all use and distribution of the Licensed Software "
        "within thirty (30) days; "
        "(c) Licensee shall return or destroy all copies of the Licensed Software "
        "and certify such destruction in writing; and "
        "(d) all accrued payment obligations shall survive termination."
    )

    # 8. Warranties
    doc.add_heading("8. Warranties and Disclaimers", level=1)
    doc.add_paragraph(
        "8.1 Licensor warrants that: "
        "(a) the Licensed Software will perform substantially in accordance with "
        "its documentation for a period of twelve (12) months from delivery; "
        "(b) Licensor has the right to grant the licenses set forth herein; and "
        "(c) to Licensor's knowledge, the Licensed Software does not infringe any "
        "third-party Intellectual Property Rights."
    )
    doc.add_paragraph(
        "8.2 EXCEPT AS EXPRESSLY SET FORTH IN SECTION 8.1, THE LICENSED SOFTWARE "
        "IS PROVIDED \"AS IS\" AND LICENSOR DISCLAIMS ALL OTHER WARRANTIES, EXPRESS "
        "OR IMPLIED, INCLUDING WARRANTIES OF MERCHANTABILITY, FITNESS FOR A "
        "PARTICULAR PURPOSE, AND NON-INFRINGEMENT."
    )

    # 9. Limitation of Liability
    doc.add_heading("9. Limitation of Liability", level=1)
    doc.add_paragraph(
        "9.1 Cap on Liability. THE TOTAL AGGREGATE LIABILITY OF EITHER PARTY UNDER "
        "THIS AGREEMENT SHALL NOT EXCEED THE GREATER OF: (A) THE TOTAL FEES PAID "
        "OR PAYABLE BY LICENSEE DURING THE TWELVE (12) MONTH PERIOD IMMEDIATELY "
        "PRECEDING THE EVENT GIVING RISE TO THE CLAIM; OR (B) EUR 5,000,000 "
        "(FIVE MILLION EUROS)."
    )
    doc.add_paragraph(
        "9.2 Exclusion of Consequential Damages. IN NO EVENT SHALL EITHER PARTY "
        "BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, CONSEQUENTIAL, OR "
        "PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, REVENUE, DATA, OR BUSINESS "
        "OPPORTUNITY, REGARDLESS OF THE THEORY OF LIABILITY."
    )
    doc.add_paragraph(
        "9.3 Exceptions. The limitations in Sections 9.1 and 9.2 shall not apply "
        "to: (a) breaches of Section 2 (License Grant) or Section 3 (Intellectual "
        "Property); (b) indemnification obligations under Section 10; or "
        "(c) willful misconduct or gross negligence."
    )

    # 10. Indemnification
    doc.add_heading("10. Indemnification", level=1)
    doc.add_paragraph(
        "10.1 Licensor shall indemnify, defend, and hold harmless Licensee from "
        "any third-party claims alleging that the Licensed Software infringes any "
        "patent, copyright, or trade secret, provided that Licensee promptly "
        "notifies Licensor of such claim and cooperates in the defense."
    )
    doc.add_paragraph(
        "10.2 Licensee shall indemnify, defend, and hold harmless Licensor from "
        "any claims arising from: (a) Licensee's distribution of the Licensed "
        "Software; (b) Licensee's Derivative Works; or (c) Licensee's breach of "
        "this Agreement."
    )

    # 11. Insurance
    doc.add_heading("11. Insurance", level=1)
    doc.add_paragraph(
        "11.1 During the term of this Agreement, Licensee shall maintain the "
        "following insurance coverage: "
        "(a) Commercial General Liability insurance with minimum limits of "
        "EUR 5,000,000 per occurrence and EUR 10,000,000 in the aggregate; "
        "(b) Professional Liability (Errors & Omissions) insurance with minimum "
        "limits of EUR 3,000,000 per claim and EUR 6,000,000 in the aggregate; "
        "(c) Cyber Liability insurance with minimum limits of EUR 2,000,000 per "
        "occurrence; and "
        "(d) Workers' Compensation insurance as required by applicable law."
    )
    doc.add_paragraph(
        "11.2 Licensee shall provide Licensor with certificates of insurance "
        "evidencing the required coverage within thirty (30) days of the Effective "
        "Date and upon each renewal thereof. Licensor shall be named as an "
        "additional insured on the Commercial General Liability policy."
    )

    # 12. Governing Law
    doc.add_heading("12. Governing Law and Jurisdiction", level=1)
    doc.add_paragraph(
        "12.1 This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of California, United States of America, without "
        "regard to its conflict of laws provisions."
    )
    doc.add_paragraph(
        "12.2 Any dispute arising out of or relating to this Agreement shall be "
        "submitted to the exclusive jurisdiction of the state and federal courts "
        "located in Santa Clara County, California. Each Party hereby consents to "
        "the personal jurisdiction of such courts."
    )
    doc.add_paragraph(
        "12.3 Notwithstanding Section 12.2, either Party may seek injunctive or "
        "other equitable relief in any court of competent jurisdiction to protect "
        "its Intellectual Property Rights."
    )

    # 13. General
    doc.add_heading("13. General Provisions", level=1)
    doc.add_paragraph(
        "13.1 Entire Agreement. This Agreement, together with all Exhibits attached "
        "hereto, constitutes the entire agreement between the Parties."
    )
    doc.add_paragraph(
        "13.2 Force Majeure. Neither Party shall be liable for any failure or delay "
        "in performance due to causes beyond its reasonable control, including "
        "natural disasters, war, terrorism, pandemic, government action, or failure "
        "of telecommunications infrastructure."
    )
    doc.add_paragraph(
        "13.3 Assignment. Licensee may not assign this Agreement without Licensor's "
        "prior written consent, except in connection with a merger or acquisition "
        "where the assignee assumes all obligations hereunder."
    )
    doc.add_paragraph(
        "13.4 Survival. Sections 3 (Intellectual Property), 6 (Non-Competition), "
        "8 (Warranties), 9 (Limitation of Liability), 10 (Indemnification), "
        "11 (Insurance), and 12 (Governing Law) shall survive termination or "
        "expiration of this Agreement."
    )
    doc.add_paragraph(
        "13.5 Export Compliance. Licensee shall comply with all applicable export "
        "control laws and regulations, including the U.S. Export Administration "
        "Regulations (EAR) and EU dual-use regulations."
    )

    # Exhibits reference
    doc.add_heading("EXHIBITS", level=1)

    # Exhibit A - Licensed Software Specifications
    doc.add_heading("Exhibit A — Licensed Software Specifications", level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    headers = ["Module", "Version", "License Type"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ("Core Platform", "12.4.1", "Full"),
        ("Analytics Engine", "12.4.1", "Full"),
        ("API Gateway", "12.3.0", "Full"),
        ("Mobile SDK", "12.2.5", "Restricted"),
        ("AI/ML Module", "12.1.0", "Evaluation Only"),
    ]
    for r, (mod, ver, lic) in enumerate(data, 1):
        table.cell(r, 0).text = mod
        table.cell(r, 1).text = ver
        table.cell(r, 2).text = lic

    # Exhibit B - Fee Schedule
    doc.add_heading("Exhibit B — Fee Schedule", level=2)
    fee_table = doc.add_table(rows=5, cols=3)
    fee_table.style = "Table Grid"
    for i, h in enumerate(["Item", "Amount (EUR)", "Frequency"]):
        fee_table.cell(0, i).text = h
    fees = [
        ("Annual License Fee", "2,400,000", "Annual (quarterly installments)"),
        ("Royalty Rate", "15%", "Quarterly"),
        ("Minimum Commitment", "1,200,000", "Annual"),
        ("Early Termination Fee", "1,200,000", "One-time"),
    ]
    for r, (item, amt, freq) in enumerate(fees, 1):
        fee_table.cell(r, 0).text = item
        fee_table.cell(r, 1).text = amt
        fee_table.cell(r, 2).text = freq

    # Signature
    doc.add_paragraph("")
    doc.add_paragraph(
        "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the "
        "Effective Date."
    )
    sig = doc.add_table(rows=4, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.cell(0, 0).text = "APEX SOFTWARE CORPORATION"
    sig.cell(0, 1).text = "GLOBALTECH SOLUTIONS GMBH"
    sig.cell(1, 0).text = "By: ________________________"
    sig.cell(1, 1).text = "By: ________________________"
    sig.cell(2, 0).text = "Name: Michael Torres, VP Licensing"
    sig.cell(2, 1).text = "Name: Dr. Klaus Weber, Geschäftsführer"
    sig.cell(3, 0).text = "Date: January 1, 2025"
    sig.cell(3, 1).text = "Date: January 1, 2025"

    out = FIXTURES_DIR / "cuad_license_agreement.docx"
    doc.save(str(out))
    print(f"Created: {out}  ({out.stat().st_size:,} bytes)")
    return out


# ─────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 60)
    print("Generating LegalBench / CUAD benchmark fixtures")
    print("=" * 60)
    create_legalbench_nda()
    create_cuad_license_agreement()
    print("\nDone! Fixtures ready for Postman integration tests.")
