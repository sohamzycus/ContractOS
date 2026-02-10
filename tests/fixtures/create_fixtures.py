"""Generate test fixture documents for ContractOS tests.

Run this script once to create .docx and .pdf test fixtures:
    python tests/fixtures/create_fixtures.py
"""

from __future__ import annotations

from pathlib import Path

FIXTURES_DIR = Path(__file__).parent


def create_simple_procurement_docx() -> None:
    """Create a simple procurement contract .docx for testing."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    doc.add_heading("Master Services Agreement", level=0)

    # Section 1: Definitions
    doc.add_heading("1. Definitions", level=1)
    doc.add_paragraph(
        'This Master Services Agreement ("Agreement") is entered into between '
        'Alpha Corp, hereinafter referred to as "Buyer", and Beta Services Ltd, '
        'hereinafter referred to as "Vendor".'
    )
    doc.add_paragraph(
        '"Effective Date" shall mean January 1, 2025.'
    )
    doc.add_paragraph(
        '"Service Period" shall mean the period of thirty (30) days from the Effective Date.'
    )

    # Section 2: Scope of Services
    doc.add_heading("2. Scope of Services", level=1)
    doc.add_paragraph(
        "The Vendor shall provide IT maintenance services for all equipment "
        "listed in Schedule A at the locations specified in Schedule B."
    )

    # Section 3: Payment Terms
    doc.add_heading("3. Payment Terms", level=1)
    doc.add_paragraph(
        "The Buyer shall pay the Vendor a total fee of $150,000.00 (One Hundred "
        "Fifty Thousand US Dollars) for the services described herein."
    )
    doc.add_paragraph(
        "Payment shall be made Net 90 from invoice date."
    )

    # Section 4: Termination
    doc.add_heading("4. Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this Agreement by providing sixty (60) days "
        "written notice to the other party. Subject to the notice period as "
        "mentioned in Section 3.2.1, the termination shall be effective upon "
        "expiry of the notice period."
    )
    doc.add_paragraph(
        "Termination may occur for the following reasons: material breach, "
        "insolvency, or mutual agreement."
    )

    # Section 5: Confidentiality
    doc.add_heading("5. Confidentiality", level=1)
    doc.add_paragraph(
        "Both parties agree to maintain the confidentiality of all proprietary "
        "information exchanged during the term of this Agreement, as further "
        "detailed in Appendix A."
    )

    # Table: Schedule A — Products
    doc.add_heading("Schedule A: Products", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Product", "Category", "Quantity"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ("Dell Inspiron 15", "IT Equipment", "50"),
        ("HP LaserJet Pro", "Office Equipment", "25"),
        ("Cisco Router 4000", "Network Equipment", "10"),
    ]
    for row_idx, (prod, cat, qty) in enumerate(data, start=1):
        table.rows[row_idx].cells[0].text = prod
        table.rows[row_idx].cells[1].text = cat
        table.rows[row_idx].cells[2].text = qty

    # Table: Schedule B — Locations
    doc.add_heading("Schedule B: Locations", level=2)
    table2 = doc.add_table(rows=4, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Location"
    table2.rows[0].cells[1].text = "Region"
    locs = [("Bangalore", "India"), ("Pune", "India"), ("Mumbai", "India")]
    for row_idx, (loc, region) in enumerate(locs, start=1):
        table2.rows[row_idx].cells[0].text = loc
        table2.rows[row_idx].cells[1].text = region

    doc.save(FIXTURES_DIR / "simple_procurement.docx")
    print("Created simple_procurement.docx")


def create_simple_nda_pdf() -> None:
    """Create a simple NDA .pdf for testing using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
    except ImportError:
        print("reportlab not installed — skipping PDF fixture creation")
        return

    pdf_path = FIXTURES_DIR / "simple_nda.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Non-Disclosure Agreement", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph("1. Parties", styles["Heading2"]))
    story.append(Paragraph(
        'This Non-Disclosure Agreement ("NDA") is entered into between '
        'Gamma Inc (the "Discloser") and Delta LLC (the "Recipient").',
        styles["Normal"],
    ))
    story.append(Spacer(1, 6))

    story.append(Paragraph("2. Confidential Information", styles["Heading2"]))
    story.append(Paragraph(
        "Confidential Information includes all trade secrets, business plans, "
        "financial data, and technical specifications disclosed by the Discloser.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 6))

    story.append(Paragraph("3. Obligations", styles["Heading2"]))
    story.append(Paragraph(
        "The Recipient agrees to hold all Confidential Information in strict "
        "confidence for a period of twenty-four (24) months from the date of disclosure.",
        styles["Normal"],
    ))
    story.append(Spacer(1, 6))

    story.append(Paragraph("4. Termination", styles["Heading2"]))
    story.append(Paragraph(
        "This NDA may be terminated by either party with thirty (30) days written notice. "
        "Upon termination, the Recipient shall return all Confidential Information "
        "as specified in Section 2.",
        styles["Normal"],
    ))

    # Simple table
    story.append(Spacer(1, 12))
    story.append(Paragraph("Schedule: Key Dates", styles["Heading3"]))
    table_data = [
        ["Event", "Date"],
        ["Effective Date", "January 1, 2025"],
        ["Expiry Date", "December 31, 2026"],
    ]
    t = Table(table_data)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
    ]))
    story.append(t)

    doc.build(story)
    print("Created simple_nda.pdf")


if __name__ == "__main__":
    create_simple_procurement_docx()
    create_simple_nda_pdf()
