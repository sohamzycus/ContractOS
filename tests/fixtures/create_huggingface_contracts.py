"""Create realistic contract documents based on HuggingFace legal datasets.

Sources:
- ContractNLI (Stanford NLP): 607 annotated NDAs — https://huggingface.co/datasets/kiddothe2b/contract-nli
- CUAD (Atticus Project): 510 real commercial contracts — https://huggingface.co/datasets/theatticusproject/cuad

These fixtures simulate real-world procurement contracts with the clause structures,
entity bindings, and cross-references found in the CUAD and ContractNLI datasets.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt

FIXTURES_DIR = Path(__file__).parent


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.size = Pt(14 if level == 1 else 12)


def create_master_services_agreement() -> str:
    """Master Services Agreement — based on CUAD commercial contract patterns.

    Covers: payment, termination, IP, indemnification, limitation of liability,
    insurance, non-compete, change of control, audit rights.
    """
    doc = Document()

    doc.add_heading("MASTER SERVICES AGREEMENT", 0)
    doc.add_paragraph(
        "This Master Services Agreement (the \"Agreement\") is entered into as of January 15, 2025 "
        "(the \"Effective Date\") by and between:"
    )
    doc.add_paragraph(
        "GlobalTech Solutions Inc., a Delaware corporation with principal offices at "
        "500 Innovation Drive, San Jose, CA 95134 (\"Service Provider\" or \"GlobalTech\"), and"
    )
    doc.add_paragraph(
        "Meridian Manufacturing Corp., a New York corporation with principal offices at "
        "1200 Industrial Boulevard, Rochester, NY 14623 (\"Client\" or \"Meridian\")."
    )
    doc.add_paragraph(
        "Service Provider and Client are each referred to herein as a \"Party\" and "
        "collectively as the \"Parties\"."
    )

    _add_heading(doc, "1. DEFINITIONS")
    doc.add_paragraph(
        "\"Confidential Information\" means all non-public information disclosed by either Party "
        "to the other Party, whether orally, in writing, or electronically, that is designated as "
        "confidential or that reasonably should be understood to be confidential given the nature "
        "of the information and the circumstances of disclosure."
    )
    doc.add_paragraph(
        "\"Deliverables\" means the work product, reports, software, documentation, and other "
        "materials to be delivered by Service Provider to Client under a Statement of Work."
    )
    doc.add_paragraph(
        "\"Intellectual Property\" or \"IP\" means all patents, copyrights, trademarks, trade secrets, "
        "and other intellectual property rights."
    )
    doc.add_paragraph(
        "\"Statement of Work\" or \"SOW\" means a document executed by both Parties that describes "
        "the specific services, deliverables, timelines, and fees for a particular engagement."
    )

    _add_heading(doc, "2. SCOPE OF SERVICES")
    doc.add_paragraph(
        "2.1 Service Provider shall provide the services described in each Statement of Work "
        "(the \"Services\"). Each SOW shall be governed by the terms of this Agreement."
    )
    doc.add_paragraph(
        "2.2 Service Provider shall assign qualified personnel with appropriate skills and "
        "experience to perform the Services. Client may request replacement of any personnel "
        "who do not meet reasonable performance standards."
    )

    _add_heading(doc, "3. FEES AND PAYMENT")
    doc.add_paragraph(
        "3.1 Client shall pay Service Provider the fees set forth in each SOW. The base annual "
        "fee for ongoing managed services is $2,400,000 (Two Million Four Hundred Thousand "
        "Dollars), payable in equal monthly installments of $200,000."
    )
    doc.add_paragraph(
        "3.2 Payment terms are Net 45 days from the date of invoice. Late payments shall "
        "accrue interest at the rate of 1.5% per month or the maximum rate permitted by law, "
        "whichever is less."
    )
    doc.add_paragraph(
        "3.3 Service Provider shall submit invoices monthly in arrears. Each invoice shall "
        "include a detailed breakdown of services performed, hours worked, and expenses incurred."
    )
    doc.add_paragraph(
        "3.4 Annual fee adjustments shall not exceed 3% per year, tied to the Consumer Price "
        "Index (CPI) published by the U.S. Bureau of Labor Statistics."
    )

    _add_heading(doc, "4. TERM AND TERMINATION")
    doc.add_paragraph(
        "4.1 This Agreement shall commence on the Effective Date and continue for an initial "
        "term of three (3) years (the \"Initial Term\"), unless earlier terminated as provided herein."
    )
    doc.add_paragraph(
        "4.2 After the Initial Term, this Agreement shall automatically renew for successive "
        "one (1) year periods (each a \"Renewal Term\") unless either Party provides written "
        "notice of non-renewal at least ninety (90) days prior to the end of the then-current term."
    )
    doc.add_paragraph(
        "4.3 Either Party may terminate this Agreement for cause if the other Party materially "
        "breaches any provision of this Agreement and fails to cure such breach within thirty (30) "
        "days after receiving written notice thereof."
    )
    doc.add_paragraph(
        "4.4 Client may terminate this Agreement for convenience upon sixty (60) days prior "
        "written notice to Service Provider, subject to payment of all fees accrued through the "
        "effective date of termination and a termination fee equal to three (3) months of the "
        "then-current monthly fee."
    )

    _add_heading(doc, "5. INTELLECTUAL PROPERTY")
    doc.add_paragraph(
        "5.1 All Deliverables created by Service Provider specifically for Client under this "
        "Agreement shall be considered \"work made for hire\" and shall be the exclusive property "
        "of Client."
    )
    doc.add_paragraph(
        "5.2 Service Provider retains all rights in its pre-existing intellectual property, "
        "tools, methodologies, and know-how (\"Service Provider IP\"). Service Provider grants "
        "Client a non-exclusive, perpetual, royalty-free license to use Service Provider IP "
        "solely as incorporated in the Deliverables."
    )

    _add_heading(doc, "6. CONFIDENTIALITY")
    doc.add_paragraph(
        "6.1 Each Party agrees to hold the other Party's Confidential Information in strict "
        "confidence and not to disclose it to any third party without the prior written consent "
        "of the disclosing Party."
    )
    doc.add_paragraph(
        "6.2 The obligations of confidentiality shall survive the termination of this Agreement "
        "for a period of five (5) years."
    )
    doc.add_paragraph(
        "6.3 Confidential Information does not include information that: (a) is or becomes "
        "publicly available through no fault of the receiving Party; (b) was known to the "
        "receiving Party prior to disclosure; (c) is independently developed by the receiving "
        "Party; or (d) is disclosed pursuant to a court order or governmental requirement."
    )

    _add_heading(doc, "7. INDEMNIFICATION")
    doc.add_paragraph(
        "7.1 Service Provider shall indemnify, defend, and hold harmless Client and its "
        "officers, directors, employees, and agents from and against any and all claims, damages, "
        "losses, liabilities, costs, and expenses (including reasonable attorneys' fees) arising "
        "from: (a) Service Provider's breach of this Agreement; (b) Service Provider's negligence "
        "or willful misconduct; (c) any claim that the Deliverables infringe any third-party "
        "intellectual property rights."
    )
    doc.add_paragraph(
        "7.2 Client shall indemnify Service Provider against claims arising from Client's "
        "use of the Deliverables in a manner not authorized by this Agreement."
    )

    _add_heading(doc, "8. LIMITATION OF LIABILITY")
    doc.add_paragraph(
        "8.1 EXCEPT FOR INDEMNIFICATION OBLIGATIONS AND BREACHES OF CONFIDENTIALITY, NEITHER "
        "PARTY'S AGGREGATE LIABILITY UNDER THIS AGREEMENT SHALL EXCEED THE TOTAL FEES PAID OR "
        "PAYABLE BY CLIENT IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM, WHICH IS ESTIMATED "
        "AT $2,400,000."
    )
    doc.add_paragraph(
        "8.2 IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, SPECIAL, "
        "CONSEQUENTIAL, OR PUNITIVE DAMAGES, INCLUDING LOSS OF PROFITS, DATA, OR BUSINESS "
        "OPPORTUNITIES, REGARDLESS OF WHETHER SUCH DAMAGES WERE FORESEEABLE."
    )

    _add_heading(doc, "9. INSURANCE")
    doc.add_paragraph(
        "9.1 Service Provider shall maintain the following insurance coverage throughout the "
        "term of this Agreement: (a) Commercial General Liability insurance with limits of not "
        "less than $5,000,000 per occurrence and $10,000,000 in the aggregate; (b) Professional "
        "Liability (Errors & Omissions) insurance with limits of not less than $5,000,000 per "
        "claim; (c) Workers' Compensation insurance as required by applicable law; (d) Cyber "
        "Liability insurance with limits of not less than $3,000,000."
    )

    _add_heading(doc, "10. NON-SOLICITATION")
    doc.add_paragraph(
        "10.1 During the term of this Agreement and for a period of twelve (12) months following "
        "termination, neither Party shall directly or indirectly solicit for employment any "
        "employee of the other Party who was involved in the performance of Services under this "
        "Agreement, without the prior written consent of the other Party."
    )

    _add_heading(doc, "11. AUDIT RIGHTS")
    doc.add_paragraph(
        "11.1 Client shall have the right, upon thirty (30) days prior written notice, to audit "
        "Service Provider's records, systems, and facilities related to the Services, no more "
        "than once per calendar year. Service Provider shall cooperate fully with any such audit."
    )

    _add_heading(doc, "12. GOVERNING LAW AND DISPUTE RESOLUTION")
    doc.add_paragraph(
        "12.1 This Agreement shall be governed by and construed in accordance with the laws of "
        "the State of New York, without regard to its conflict of laws principles."
    )
    doc.add_paragraph(
        "12.2 Any dispute arising out of or relating to this Agreement shall first be submitted "
        "to mediation. If mediation is unsuccessful, the dispute shall be resolved by binding "
        "arbitration administered by the American Arbitration Association (AAA) under its "
        "Commercial Arbitration Rules."
    )

    _add_heading(doc, "13. GENERAL PROVISIONS")
    doc.add_paragraph(
        "13.1 Force Majeure. Neither Party shall be liable for any failure or delay in "
        "performance due to circumstances beyond its reasonable control, including but not "
        "limited to acts of God, war, terrorism, pandemic, government actions, or natural disasters."
    )
    doc.add_paragraph(
        "13.2 Assignment. Neither Party may assign this Agreement without the prior written "
        "consent of the other Party, except in connection with a merger, acquisition, or sale "
        "of substantially all of its assets."
    )
    doc.add_paragraph(
        "13.3 Entire Agreement. This Agreement, together with all SOWs and exhibits, constitutes "
        "the entire agreement between the Parties and supersedes all prior negotiations, "
        "representations, and agreements."
    )

    path = str(FIXTURES_DIR / "hf_master_services_agreement.docx")
    doc.save(path)
    return path


def create_software_license_agreement() -> str:
    """Software License Agreement — based on CUAD SaaS/license contract patterns.

    Covers: license grant, SLA, data processing, warranty, limitation of liability.
    """
    doc = Document()

    doc.add_heading("SOFTWARE LICENSE AND SUBSCRIPTION AGREEMENT", 0)
    doc.add_paragraph(
        "This Software License and Subscription Agreement (\"Agreement\") is effective as of "
        "March 1, 2025 (\"Effective Date\") between:"
    )
    doc.add_paragraph(
        "CloudVault Technologies Ltd., a company incorporated under the laws of England and Wales, "
        "with registered offices at 45 Canary Wharf, London E14 5AB, United Kingdom "
        "(\"Licensor\" or \"CloudVault\"), and"
    )
    doc.add_paragraph(
        "Pacific Rim Logistics Inc., a California corporation with principal offices at "
        "8800 Harbor Drive, Long Beach, CA 90802 (\"Licensee\" or \"Pacific Rim\")."
    )

    _add_heading(doc, "1. LICENSE GRANT")
    doc.add_paragraph(
        "1.1 Subject to the terms of this Agreement, Licensor grants Licensee a non-exclusive, "
        "non-transferable, worldwide license to access and use the CloudVault Enterprise Platform "
        "(the \"Software\") during the Subscription Term for up to 500 authorized users."
    )
    doc.add_paragraph(
        "1.2 The license includes access to all standard modules: Inventory Management, "
        "Shipment Tracking, Customs Compliance, Warehouse Optimization, and Analytics Dashboard."
    )
    doc.add_paragraph(
        "1.3 Licensee shall not: (a) sublicense, sell, or transfer the Software; (b) reverse "
        "engineer, decompile, or disassemble the Software; (c) use the Software to develop "
        "a competing product; (d) exceed the authorized user limit without purchasing additional licenses."
    )

    _add_heading(doc, "2. SUBSCRIPTION FEES")
    doc.add_paragraph(
        "2.1 The annual subscription fee is $750,000 (Seven Hundred Fifty Thousand Dollars) "
        "for up to 500 users. Additional users may be added at $1,500 per user per year."
    )
    doc.add_paragraph(
        "2.2 Implementation and onboarding services: a one-time fee of $150,000."
    )
    doc.add_paragraph(
        "2.3 Premium support (24/7 with 1-hour response SLA): $120,000 per year."
    )
    doc.add_paragraph(
        "2.4 All fees are payable annually in advance within thirty (30) days of invoice. "
        "Late payments accrue interest at 1% per month."
    )

    _add_heading(doc, "3. SERVICE LEVEL AGREEMENT")
    doc.add_paragraph(
        "3.1 Licensor guarantees 99.95% uptime for the Software, measured monthly, excluding "
        "scheduled maintenance windows (Sundays 2:00-6:00 AM UTC)."
    )
    doc.add_paragraph(
        "3.2 If uptime falls below 99.95% in any calendar month, Licensee shall receive "
        "service credits: (a) 99.0%-99.95%: 10% credit of monthly fee; (b) 95.0%-99.0%: "
        "25% credit; (c) below 95.0%: 50% credit."
    )
    doc.add_paragraph(
        "3.3 Response times for support tickets: Critical (P1): 1 hour; High (P2): 4 hours; "
        "Medium (P3): 8 business hours; Low (P4): 2 business days."
    )

    _add_heading(doc, "4. DATA PROCESSING AND SECURITY")
    doc.add_paragraph(
        "4.1 Licensor shall process Licensee's data solely for the purpose of providing the "
        "Software services. Licensor shall comply with all applicable data protection laws, "
        "including GDPR, CCPA, and any successor regulations."
    )
    doc.add_paragraph(
        "4.2 All data shall be encrypted at rest (AES-256) and in transit (TLS 1.3). "
        "Licensor maintains SOC 2 Type II and ISO 27001 certifications."
    )
    doc.add_paragraph(
        "4.3 Data residency: Licensee's data shall be stored exclusively in data centers "
        "located within the United States and the European Union, as selected by Licensee."
    )
    doc.add_paragraph(
        "4.4 Upon termination, Licensor shall export all Licensee data in a standard format "
        "(CSV, JSON, or SQL dump) within thirty (30) days and permanently delete all copies "
        "within sixty (60) days."
    )

    _add_heading(doc, "5. TERM AND TERMINATION")
    doc.add_paragraph(
        "5.1 The initial subscription term is two (2) years from the Effective Date. "
        "The Agreement automatically renews for successive one (1) year terms unless either "
        "Party provides sixty (60) days written notice of non-renewal."
    )
    doc.add_paragraph(
        "5.2 Either Party may terminate for material breach with thirty (30) days written "
        "notice and opportunity to cure."
    )
    doc.add_paragraph(
        "5.3 Licensee may terminate for convenience with ninety (90) days notice, subject "
        "to payment of fees for the remainder of the then-current term."
    )

    _add_heading(doc, "6. WARRANTIES")
    doc.add_paragraph(
        "6.1 Licensor warrants that the Software will perform substantially in accordance "
        "with its published documentation during the Subscription Term."
    )
    doc.add_paragraph(
        "6.2 Licensor warrants that it has the right to grant the licenses contemplated "
        "by this Agreement and that the Software does not infringe any third-party intellectual "
        "property rights."
    )
    doc.add_paragraph(
        "6.3 EXCEPT AS EXPRESSLY SET FORTH HEREIN, THE SOFTWARE IS PROVIDED \"AS IS\" AND "
        "LICENSOR DISCLAIMS ALL OTHER WARRANTIES, EXPRESS OR IMPLIED, INCLUDING WARRANTIES "
        "OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE."
    )

    _add_heading(doc, "7. LIMITATION OF LIABILITY")
    doc.add_paragraph(
        "7.1 LICENSOR'S TOTAL AGGREGATE LIABILITY SHALL NOT EXCEED THE FEES PAID BY LICENSEE "
        "IN THE TWELVE (12) MONTHS PRECEDING THE CLAIM."
    )
    doc.add_paragraph(
        "7.2 NEITHER PARTY SHALL BE LIABLE FOR INDIRECT, INCIDENTAL, SPECIAL, OR CONSEQUENTIAL "
        "DAMAGES, EXCEPT IN CASES OF WILLFUL MISCONDUCT OR BREACH OF CONFIDENTIALITY."
    )

    _add_heading(doc, "8. GOVERNING LAW")
    doc.add_paragraph(
        "8.1 This Agreement shall be governed by the laws of the State of California. "
        "Any disputes shall be resolved by arbitration in San Francisco, California under "
        "JAMS rules."
    )

    path = str(FIXTURES_DIR / "hf_software_license_agreement.docx")
    doc.save(path)
    return path


def create_supply_chain_agreement() -> str:
    """Supply Chain and Procurement Agreement — based on CUAD procurement patterns.

    Covers: supply obligations, pricing, quality, delivery, warranty, penalties.
    """
    doc = Document()

    doc.add_heading("SUPPLY CHAIN AND PROCUREMENT AGREEMENT", 0)
    doc.add_paragraph(
        "This Supply Chain and Procurement Agreement (\"Agreement\") is entered into as of "
        "February 10, 2025 by and between:"
    )
    doc.add_paragraph(
        "Apex Industrial Components GmbH, a company organized under the laws of Germany, "
        "with headquarters at Industriestraße 42, 70565 Stuttgart, Germany "
        "(\"Supplier\" or \"Apex\"), and"
    )
    doc.add_paragraph(
        "NovaTech Automotive Ltd., a company incorporated in Japan with principal offices at "
        "2-1-1 Marunouchi, Chiyoda-ku, Tokyo 100-0005, Japan (\"Buyer\" or \"NovaTech\")."
    )

    _add_heading(doc, "1. SUPPLY OBLIGATIONS")
    doc.add_paragraph(
        "1.1 Supplier shall manufacture and deliver the products specified in Exhibit A "
        "(the \"Products\") in accordance with the specifications, quantities, and delivery "
        "schedules set forth therein."
    )
    doc.add_paragraph(
        "1.2 Supplier shall maintain a minimum safety stock of 10,000 units of each Product "
        "at all times to ensure uninterrupted supply."
    )
    doc.add_paragraph(
        "1.3 Supplier shall provide Buyer with a rolling 12-month demand forecast and shall "
        "adjust production capacity accordingly within thirty (30) days of receiving updated "
        "forecasts from Buyer."
    )

    _add_heading(doc, "2. PRICING AND PAYMENT")
    doc.add_paragraph(
        "2.1 Unit prices for the Products are set forth in Exhibit B. The total estimated "
        "annual contract value is EUR 18,500,000 (Eighteen Million Five Hundred Thousand Euros) "
        "based on projected volumes."
    )
    doc.add_paragraph(
        "2.2 Prices are fixed for the first two (2) years. Thereafter, price adjustments "
        "shall not exceed 2.5% per year and must be supported by documented cost increases "
        "in raw materials or labor."
    )
    doc.add_paragraph(
        "2.3 Payment terms: Net 60 days from date of delivery and acceptance. Buyer may "
        "withhold payment for non-conforming Products until defects are remedied."
    )
    doc.add_paragraph(
        "2.4 Volume discounts: Orders exceeding 50,000 units per quarter shall receive a "
        "5% discount. Orders exceeding 100,000 units shall receive a 10% discount."
    )

    _add_heading(doc, "3. QUALITY REQUIREMENTS")
    doc.add_paragraph(
        "3.1 All Products shall conform to the specifications in Exhibit A and shall comply "
        "with ISO 9001:2015, IATF 16949, and all applicable automotive industry standards."
    )
    doc.add_paragraph(
        "3.2 The acceptable defect rate shall not exceed 50 parts per million (PPM). If the "
        "defect rate exceeds 100 PPM in any quarter, Buyer may require Supplier to implement "
        "a corrective action plan within fifteen (15) days."
    )
    doc.add_paragraph(
        "3.3 Supplier shall maintain full traceability of all raw materials and components "
        "used in the Products. Buyer shall have the right to audit Supplier's quality management "
        "system upon thirty (30) days written notice."
    )

    _add_heading(doc, "4. DELIVERY AND LOGISTICS")
    doc.add_paragraph(
        "4.1 Delivery terms: DDP (Delivered Duty Paid) to Buyer's designated facilities in "
        "Tokyo, Japan and Detroit, Michigan, USA (Incoterms 2020)."
    )
    doc.add_paragraph(
        "4.2 On-time delivery rate shall be at least 98%. Supplier shall notify Buyer of any "
        "anticipated delivery delays at least fourteen (14) days in advance."
    )
    doc.add_paragraph(
        "4.3 Packaging shall comply with Buyer's packaging specifications and all applicable "
        "transportation regulations. Supplier shall use recyclable packaging materials where feasible."
    )

    _add_heading(doc, "5. WARRANTIES")
    doc.add_paragraph(
        "5.1 Supplier warrants that all Products shall be free from defects in materials and "
        "workmanship for a period of thirty-six (36) months from the date of delivery or "
        "24 months from installation in Buyer's end products, whichever occurs first."
    )
    doc.add_paragraph(
        "5.2 Supplier warrants that all Products comply with applicable laws and regulations, "
        "including REACH, RoHS, and conflict minerals regulations."
    )

    _add_heading(doc, "6. PENALTIES AND REMEDIES")
    doc.add_paragraph(
        "6.1 Late delivery penalty: 1% of the value of the delayed shipment per week of delay, "
        "up to a maximum of 10% of the shipment value."
    )
    doc.add_paragraph(
        "6.2 Quality failure penalty: If the defect rate exceeds 100 PPM for two consecutive "
        "quarters, Supplier shall pay liquidated damages of EUR 500,000 per quarter until the "
        "defect rate is brought below 50 PPM."
    )
    doc.add_paragraph(
        "6.3 Recall costs: Supplier shall bear all costs associated with product recalls "
        "caused by defects in the Products, including but not limited to logistics, replacement "
        "parts, labor, and customer notifications."
    )

    _add_heading(doc, "7. TERM AND TERMINATION")
    doc.add_paragraph(
        "7.1 This Agreement shall have an initial term of five (5) years from the Effective Date."
    )
    doc.add_paragraph(
        "7.2 Either Party may terminate for material breach with sixty (60) days written notice "
        "and opportunity to cure."
    )
    doc.add_paragraph(
        "7.3 Buyer may terminate for convenience with one hundred eighty (180) days notice, "
        "subject to payment for all Products ordered and in production."
    )
    doc.add_paragraph(
        "7.4 Upon termination, Supplier shall continue to supply spare parts for a minimum "
        "of ten (10) years at the prices in effect at the time of termination, adjusted annually "
        "for inflation."
    )

    _add_heading(doc, "8. LIMITATION OF LIABILITY")
    doc.add_paragraph(
        "8.1 Supplier's total liability shall not exceed 150% of the annual contract value "
        "(EUR 27,750,000), except for claims arising from product recalls, willful misconduct, "
        "or breach of confidentiality."
    )

    _add_heading(doc, "9. GOVERNING LAW")
    doc.add_paragraph(
        "9.1 This Agreement shall be governed by the United Nations Convention on Contracts "
        "for the International Sale of Goods (CISG) and, to the extent not covered by CISG, "
        "by the laws of Germany."
    )
    doc.add_paragraph(
        "9.2 Disputes shall be resolved by arbitration under ICC rules, with the seat of "
        "arbitration in Zurich, Switzerland. The language of arbitration shall be English."
    )

    path = str(FIXTURES_DIR / "hf_supply_chain_agreement.docx")
    doc.save(path)
    return path


if __name__ == "__main__":
    print("Creating HuggingFace-sourced contract fixtures...")
    p1 = create_master_services_agreement()
    print(f"  Created: {p1}")
    p2 = create_software_license_agreement()
    print(f"  Created: {p2}")
    p3 = create_supply_chain_agreement()
    print(f"  Created: {p3}")
    print("Done. 3 realistic procurement contracts created.")
