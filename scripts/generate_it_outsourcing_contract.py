"""Generate the complex IT outsourcing contract (DOCX) for demo samples."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def build_contract():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IT OUTSOURCING SERVICES AGREEMENT")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Contract Reference: ITOSA-2024-0847")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    add_para(doc, 'This IT Outsourcing Services Agreement (the "Agreement") is entered into as of January 15, 2024 (the "Effective Date") by and between:')
    doc.add_paragraph()

    add_para(doc, 'Meridian Global Holdings, Inc., a corporation organized and existing under the laws of the State of Delaware, with its principal place of business at 500 Park Avenue, Suite 3200, New York, NY 10022, United States (hereinafter referred to as the "Client" or "Meridian"); and')
    doc.add_paragraph()

    add_para(doc, 'TechServe Solutions Private Limited, a company incorporated under the laws of India, with its registered office at Tower B, Cyber City, DLF Phase 2, Gurugram, Haryana 122002, India (hereinafter referred to as the "Service Provider", "Vendor", or "TechServe").')
    doc.add_paragraph()

    add_para(doc, 'WHEREAS, the Client desires to outsource certain information technology services and the Service Provider has the expertise, resources, and capability to provide such services;')
    add_para(doc, 'WHEREAS, the Parties wish to establish the terms and conditions under which the Service Provider shall deliver IT outsourcing services to the Client across multiple geographies;')
    add_para(doc, 'NOW, THEREFORE, in consideration of the mutual covenants and agreements set forth herein, and for other good and valuable consideration, the receipt and sufficiency of which are hereby acknowledged, the Parties agree as follows:')

    doc.add_paragraph()

    # Section 1: Definitions
    add_heading(doc, "1. Definitions and Interpretation", level=1)

    add_para(doc, 'In this Agreement, unless the context otherwise requires, the following terms shall have the meanings ascribed to them below:')
    doc.add_paragraph()

    definitions = [
        ('"Authorized Users"', 'means employees, contractors, and agents of the Client who are authorized to access the Services.'),
        ('"Base Fee"', 'means the annual base service fee of USD 9,500,000 (Nine Million Five Hundred Thousand United States Dollars) payable by the Client to the Service Provider, as set out in Schedule 2 (Pricing).'),
        ('"Business Day"', 'means any day other than a Saturday, Sunday, or public holiday in both New York, United States and Gurugram, India.'),
        ('"Change Request"', 'means a formal request for modification to the Services, submitted in accordance with Section 14 (Change Management).'),
        ('"Confidential Information"', 'means all information disclosed by one Party to the other, whether orally, in writing, or electronically, that is designated as confidential or that reasonably should be understood to be confidential given the nature of the information and the circumstances of disclosure, including but not limited to trade secrets, business plans, financial data, customer lists, technical specifications, source code, algorithms, and proprietary methodologies.'),
        ('"Data Processing Agreement" or "DPA"', 'means the data processing agreement attached hereto as Schedule 5, governing the processing of Personal Data.'),
        ('"Deliverables"', 'means all work product, reports, documentation, software, and other materials created by the Service Provider in the course of performing the Services.'),
        ('"Force Majeure Event"', 'means any event beyond the reasonable control of the affected Party, including but not limited to acts of God, war, terrorism, pandemic, epidemic, earthquake, flood, fire, explosion, civil unrest, government sanctions, cyber-attack, widespread internet outage, or labor strike not involving the affected Party\'s employees.'),
        ('"FTE"', 'means a full-time equivalent resource dedicated to the performance of the Services, working a minimum of 40 hours per week.'),
        ('"Intellectual Property" or "IP"', 'means all patents, copyrights, trademarks, trade secrets, know-how, inventions, designs, software, databases, and all other intellectual property rights, whether registered or unregistered.'),
        ('"Key Personnel"', 'means the individuals identified in Schedule 4 who are essential to the delivery of the Services and may not be replaced without the Client\'s prior written consent.'),
        ('"Personal Data"', 'means any information relating to an identified or identifiable natural person, as defined under applicable data protection laws including GDPR, CCPA, and India\'s Digital Personal Data Protection Act, 2023.'),
        ('"Service Credits"', 'means the credits payable by the Service Provider to the Client in the event of failure to meet the Service Levels, as calculated in accordance with Schedule 3 (Service Level Agreement).'),
        ('"Service Levels" or "SLAs"', 'means the performance standards and metrics set out in Schedule 3 (Service Level Agreement).'),
        ('"Term"', 'means the initial term of five (5) years commencing on the Effective Date, unless terminated earlier in accordance with Section 12 (Termination) or extended pursuant to Section 2.2.'),
        ('"Total Contract Value"', 'means USD 47,500,000 (Forty-Seven Million Five Hundred Thousand United States Dollars) over the initial five-year Term.'),
        ('"Transition Period"', 'means the period of up to twelve (12) months following termination or expiration during which the Service Provider shall provide transition assistance services as described in Section 12.6.'),
    ]

    for term, definition in definitions:
        p = doc.add_paragraph()
        run = p.add_run(term)
        run.bold = True
        p.add_run(f" {definition}")

    doc.add_paragraph()

    # Section 2: Term and Renewal
    add_heading(doc, "2. Term and Renewal", level=1)

    add_para(doc, '2.1 Initial Term. This Agreement shall commence on the Effective Date and shall continue for an initial period of five (5) years (the "Initial Term"), expiring on January 14, 2029, unless terminated earlier in accordance with Section 12.')
    add_para(doc, '2.2 Renewal. Upon expiration of the Initial Term, this Agreement may be renewed for successive periods of two (2) years each (each a "Renewal Term"), upon mutual written agreement of the Parties, provided that written notice of intent to renew is given at least one hundred and eighty (180) days prior to the expiration of the then-current Term.')
    add_para(doc, '2.3 Ramp-Up Period. The first ninety (90) days following the Effective Date shall constitute the Ramp-Up Period, during which the Service Provider shall transition services from the Client\'s incumbent providers and achieve full operational capability. Service Level penalties shall not apply during the Ramp-Up Period.')

    # Section 3: Scope of Services
    add_heading(doc, "3. Scope of Services", level=1)

    add_para(doc, '3.1 Service Categories. The Service Provider shall provide the following categories of IT outsourcing services (collectively, the "Services"):')

    services = [
        ("a)", "Infrastructure Management Services — including server administration, network management, cloud infrastructure (AWS, Azure), storage management, and disaster recovery operations across 12 data center locations;"),
        ("b)", "Application Development and Maintenance — including custom software development, legacy application modernization, bug fixes, enhancements, and version upgrades for the Client's portfolio of 47 enterprise applications;"),
        ("c)", "Help Desk and End-User Support — providing Level 1, Level 2, and Level 3 technical support to approximately 8,500 end users across 12 geographic locations, operating 24/7/365;"),
        ("d)", "Cybersecurity Operations — including Security Operations Center (SOC) monitoring, vulnerability management, incident response, penetration testing, and compliance monitoring;"),
        ("e)", "Data Analytics and Business Intelligence — including data warehouse management, ETL pipeline operations, dashboard development, and ad-hoc reporting services;"),
        ("f)", "Cloud Migration and Transformation — including assessment, planning, and execution of the Client's cloud migration strategy targeting 85% cloud adoption by Year 3 of the Agreement."),
    ]

    for label, text in services:
        p = doc.add_paragraph()
        run = p.add_run(f"{label} ")
        run.bold = True
        p.add_run(text)

    add_para(doc, '3.2 Resource Commitment. The Service Provider shall deploy a minimum of 350 FTEs dedicated to the Client\'s account, distributed across the following locations:')

    add_table(doc,
        ["Location", "FTE Count", "Primary Function", "Operating Hours"],
        [
            ["Gurugram, India (Primary)", "180", "Application Development, Infrastructure", "24/7"],
            ["Bangalore, India", "65", "Cybersecurity SOC, Data Analytics", "24/7"],
            ["Hyderabad, India", "40", "Cloud Migration, DevOps", "IST Business Hours"],
            ["New York, USA", "25", "Client Liaison, Project Management", "EST Business Hours"],
            ["London, UK", "15", "European Operations, GDPR Compliance", "GMT Business Hours"],
            ["Singapore", "10", "APAC Support", "SGT Business Hours"],
            ["Toronto, Canada", "8", "Canadian Operations", "EST Business Hours"],
            ["Sydney, Australia", "7", "ANZ Support", "AEST Business Hours"],
        ]
    )

    add_para(doc, '3.3 Exclusions. The following are expressly excluded from the scope of Services: (a) procurement of third-party hardware or software licenses on behalf of the Client; (b) physical data center facility management; (c) telecommunications carrier management; and (d) end-user device procurement and provisioning.')

    # Section 4: Service Levels
    add_heading(doc, "4. Service Level Agreement", level=1)

    add_para(doc, '4.1 Service Level Framework. The Service Provider shall meet or exceed the Service Levels specified in Schedule 3. Service Levels are measured on a monthly basis and reported within five (5) Business Days of each calendar month end.')

    add_para(doc, '4.2 Severity Classification. Incidents shall be classified according to the following severity framework:')

    add_table(doc,
        ["Severity", "Description", "Response Time", "Resolution Time", "Escalation"],
        [
            ["Severity 1 (Critical)", "Complete service outage or data breach affecting production systems", "15 minutes", "4 hours", "Immediate VP-level notification"],
            ["Severity 2 (High)", "Major degradation affecting >25% of users or critical business function", "30 minutes", "8 hours", "Director notification within 1 hour"],
            ["Severity 3 (Medium)", "Partial service impact affecting <25% of users", "2 hours", "24 hours", "Manager notification within 4 hours"],
            ["Severity 4 (Low)", "Minor issue with workaround available", "4 hours", "72 hours", "Standard ticket queue"],
            ["Severity 5 (Informational)", "Enhancement request or cosmetic issue", "Next Business Day", "As scheduled", "Backlog prioritization"],
        ]
    )

    add_para(doc, '4.3 Service Credits. In the event the Service Provider fails to meet the Service Levels for any given month, the Client shall be entitled to Service Credits calculated as follows:')

    add_table(doc,
        ["SLA Metric", "Target", "Service Credit (per % below target)", "Monthly Credit Cap"],
        [
            ["System Availability", "99.95%", "2% of monthly Base Fee per 0.1% shortfall", "15% of monthly Base Fee"],
            ["Severity 1 Response Time", "15 minutes", "USD 5,000 per incident exceeding target", "10% of monthly Base Fee"],
            ["Severity 1 Resolution Time", "4 hours", "USD 10,000 per incident exceeding target", "15% of monthly Base Fee"],
            ["Severity 2 Response Time", "30 minutes", "USD 2,500 per incident exceeding target", "5% of monthly Base Fee"],
            ["Help Desk First-Call Resolution", "75%", "1% of monthly Base Fee per 5% shortfall", "10% of monthly Base Fee"],
            ["Customer Satisfaction Score", "4.0/5.0", "1.5% of monthly Base Fee per 0.5 shortfall", "7.5% of monthly Base Fee"],
        ]
    )

    add_para(doc, '4.4 Aggregate Service Credit Cap. The total Service Credits payable by the Service Provider in any calendar month shall not exceed twenty-five percent (25%) of the monthly Base Fee for that month. If Service Credits exceed twenty percent (20%) of the monthly Base Fee for three (3) consecutive months, the Client shall have the right to terminate this Agreement for cause pursuant to Section 12.3.')

    add_para(doc, '4.5 Earnback Mechanism. If the Service Provider meets or exceeds all Service Levels for three (3) consecutive months following a month in which Service Credits were applied, fifty percent (50%) of the Service Credits from the triggering month shall be credited back to the Service Provider.')

    # Section 5: Pricing and Payment
    add_heading(doc, "5. Pricing and Payment", level=1)

    add_para(doc, '5.1 Total Contract Value. The Total Contract Value for the initial five-year Term is USD 47,500,000 (Forty-Seven Million Five Hundred Thousand United States Dollars), structured as follows:')

    add_table(doc,
        ["Year", "Annual Base Fee", "Variable Component (est.)", "Total (est.)"],
        [
            ["Year 1 (2024)", "USD 9,500,000", "USD 500,000", "USD 10,000,000"],
            ["Year 2 (2025)", "USD 9,500,000", "USD 750,000", "USD 10,250,000"],
            ["Year 3 (2026)", "USD 9,200,000", "USD 800,000", "USD 10,000,000"],
            ["Year 4 (2027)", "USD 8,800,000", "USD 700,000", "USD 9,500,000"],
            ["Year 5 (2028)", "USD 8,500,000", "USD 750,000", "USD 9,250,000"],
            ["Total", "USD 45,500,000", "USD 3,500,000", "USD 49,000,000"],
        ]
    )

    add_para(doc, '5.2 Price Escalation. The Base Fee shall be subject to annual adjustment not exceeding the lesser of: (a) three percent (3%) per annum; or (b) the Consumer Price Index (CPI) increase for the preceding twelve-month period as published by the U.S. Bureau of Labor Statistics. Any adjustment exceeding 3% shall require mutual written agreement.')

    add_para(doc, '5.3 Payment Terms. The Client shall pay the Service Provider within forty-five (45) days of receipt of a valid invoice. Invoices shall be submitted monthly in arrears. Late payments shall accrue interest at the rate of 1.5% per month or the maximum rate permitted by applicable law, whichever is lower.')

    add_para(doc, '5.4 Disputed Invoices. The Client may dispute any invoice or portion thereof in good faith by providing written notice within thirty (30) days of receipt. Undisputed amounts shall be paid in accordance with Section 5.3. The Parties shall use commercially reasonable efforts to resolve any invoice dispute within sixty (60) days.')

    add_para(doc, '5.5 Quarterly Business Review. The Parties shall conduct quarterly business reviews to assess service delivery performance, financial reconciliation, and strategic alignment. The Service Provider shall provide a detailed financial report including actual vs. budgeted costs, FTE utilization, and Service Credit calculations.')

    # Section 6: Data Protection
    add_heading(doc, "6. Data Protection and Privacy", level=1)

    add_para(doc, '6.1 Compliance with Data Protection Laws. The Service Provider shall comply with all applicable data protection and privacy laws, including but not limited to:')

    laws = [
        "General Data Protection Regulation (EU) 2016/679 (GDPR);",
        "California Consumer Privacy Act (CCPA) and California Privacy Rights Act (CPRA);",
        "India's Digital Personal Data Protection Act, 2023 (DPDPA);",
        "Personal Information Protection and Electronic Documents Act (PIPEDA) (Canada);",
        "Privacy Act 1988 (Australia);",
        "Personal Data Protection Act 2012 (Singapore)."
    ]
    for law in laws:
        doc.add_paragraph(law, style="List Bullet")

    add_para(doc, '6.2 Data Processing Agreement. The Service Provider shall process Personal Data only in accordance with the DPA attached as Schedule 5. The Service Provider acts as a data processor on behalf of the Client (data controller) and shall not process Personal Data for any purpose other than the performance of the Services.')

    add_para(doc, '6.3 Data Breach Notification. In the event of a Personal Data breach, the Service Provider shall:')

    breach_steps = [
        "a) Notify the Client within twenty-four (24) hours of becoming aware of the breach;",
        "b) Provide a preliminary incident report within forty-eight (48) hours containing: (i) nature and scope of the breach, (ii) categories and approximate number of data subjects affected, (iii) categories and approximate number of Personal Data records affected, (iv) likely consequences of the breach, and (v) measures taken or proposed to address the breach;",
        "c) Provide a comprehensive root cause analysis within fourteen (14) calendar days;",
        "d) Implement remediation measures and provide evidence of their effectiveness within thirty (30) calendar days;",
        "e) Cooperate fully with the Client and any supervisory authority in the investigation of the breach."
    ]
    for step in breach_steps:
        doc.add_paragraph(step, style="List Bullet")

    add_para(doc, '6.4 Data Residency. All Personal Data of European data subjects shall be stored and processed exclusively within the European Economic Area (EEA) or in jurisdictions with an adequacy decision from the European Commission. Personal Data of U.S. data subjects shall be stored within the continental United States. The Service Provider shall not transfer Personal Data across borders without the Client\'s prior written consent and appropriate safeguards (Standard Contractual Clauses or Binding Corporate Rules).')

    add_para(doc, '6.5 Data Retention and Deletion. Upon termination or expiration of this Agreement, the Service Provider shall, at the Client\'s election, return or securely destroy all Personal Data within sixty (60) days and provide written certification of destruction. The Service Provider shall use NIST SP 800-88 compliant methods for data sanitization.')

    add_para(doc, '6.6 Audit Rights. The Client shall have the right to conduct data protection audits of the Service Provider\'s facilities, systems, and processes no more than twice per calendar year, with thirty (30) days\' prior written notice. The Service Provider shall provide reasonable cooperation and access to relevant records, systems, and personnel.')

    # Section 7: Confidentiality
    add_heading(doc, "7. Confidentiality", level=1)

    add_para(doc, '7.1 Obligations. Each Party (the "Receiving Party") shall: (a) hold in strict confidence all Confidential Information of the other Party (the "Disclosing Party"); (b) not disclose Confidential Information to any third party without the Disclosing Party\'s prior written consent; (c) use Confidential Information solely for the purposes of this Agreement; and (d) protect Confidential Information using the same degree of care it uses to protect its own confidential information, but in no event less than reasonable care.')

    add_para(doc, '7.2 Exceptions. The obligations of confidentiality shall not apply to information that: (a) is or becomes publicly available through no fault of the Receiving Party; (b) was known to the Receiving Party prior to disclosure; (c) is independently developed by the Receiving Party without use of Confidential Information; or (d) is rightfully received from a third party without restriction on disclosure.')

    add_para(doc, '7.3 Compelled Disclosure. If the Receiving Party is compelled by law, regulation, or legal process to disclose Confidential Information, it shall: (a) provide prompt written notice to the Disclosing Party (to the extent legally permitted); (b) cooperate with the Disclosing Party in seeking a protective order; and (c) disclose only the minimum amount of information required.')

    add_para(doc, '7.4 Survival. The obligations of confidentiality shall survive termination or expiration of this Agreement for a period of five (5) years, except with respect to trade secrets, which shall be protected indefinitely.')

    # Section 8: Intellectual Property
    add_heading(doc, "8. Intellectual Property Rights", level=1)

    add_para(doc, '8.1 Client IP. All Intellectual Property owned by or licensed to the Client prior to the Effective Date ("Client IP") shall remain the exclusive property of the Client. The Service Provider is granted a limited, non-exclusive, non-transferable license to use Client IP solely for the purpose of performing the Services during the Term.')

    add_para(doc, '8.2 Service Provider IP. All Intellectual Property owned by or licensed to the Service Provider prior to the Effective Date, including pre-existing tools, frameworks, libraries, and methodologies ("Service Provider IP"), shall remain the exclusive property of the Service Provider. The Client is granted a perpetual, irrevocable, royalty-free, non-exclusive license to use any Service Provider IP incorporated into the Deliverables.')

    add_para(doc, '8.3 Work Product. All Deliverables and work product created by the Service Provider specifically for the Client in the course of performing the Services ("Work Product") shall be the exclusive property of the Client. The Service Provider hereby assigns to the Client all right, title, and interest in and to the Work Product, including all Intellectual Property rights therein. The Service Provider shall execute all documents and take all actions reasonably necessary to perfect the Client\'s ownership.')

    add_para(doc, '8.4 Joint IP. Any Intellectual Property developed jointly by the Parties ("Joint IP") shall be jointly owned. Neither Party may license or assign its interest in Joint IP to a third party without the other Party\'s prior written consent. Each Party shall have the right to use Joint IP for its own internal business purposes without accounting to the other Party.')

    add_para(doc, '8.5 Open Source. The Service Provider shall not incorporate any open-source software into the Deliverables without the Client\'s prior written approval. The Service Provider shall maintain a complete inventory of all open-source components used in connection with the Services, including license types and any copyleft obligations.')

    # Section 9: Liability
    add_heading(doc, "9. Limitation of Liability", level=1)

    add_para(doc, '9.1 General Cap. Subject to Section 9.2, the total aggregate liability of either Party under or in connection with this Agreement, whether in contract, tort (including negligence), breach of statutory duty, or otherwise, shall not exceed one hundred and fifty percent (150%) of the Base Fee paid or payable in the twelve (12) months preceding the event giving rise to the claim (the "Liability Cap"). As of the Effective Date, the Liability Cap is USD 14,250,000.')

    add_para(doc, '9.2 Unlimited Liability. The Liability Cap shall not apply to:')

    unlimited = [
        "a) Liability arising from a breach of Section 6 (Data Protection and Privacy), including any Personal Data breach;",
        "b) Liability arising from a breach of Section 7 (Confidentiality);",
        "c) Liability arising from infringement of Intellectual Property rights;",
        "d) Liability arising from the Service Provider's willful misconduct or gross negligence;",
        "e) Liability arising from the Service Provider's failure to comply with applicable laws or regulations;",
        "f) Indemnification obligations under Section 10."
    ]
    for item in unlimited:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, '9.3 Consequential Damages. Neither Party shall be liable to the other for any indirect, incidental, special, consequential, or punitive damages, including but not limited to loss of profits, loss of revenue, loss of data, or loss of business opportunity, except in cases of: (a) breach of confidentiality obligations; (b) Personal Data breach; (c) willful misconduct; or (d) infringement of Intellectual Property rights.')

    add_para(doc, '9.4 Mitigation. Each Party shall take all reasonable steps to mitigate any loss or damage for which the other Party may be liable under this Agreement.')

    # Section 10: Indemnification
    add_heading(doc, "10. Indemnification", level=1)

    add_para(doc, '10.1 Service Provider Indemnification. The Service Provider shall indemnify, defend, and hold harmless the Client and its officers, directors, employees, and agents from and against any and all claims, damages, losses, liabilities, costs, and expenses (including reasonable attorneys\' fees) arising out of or relating to:')

    sp_indemnity = [
        "a) Any breach of the Service Provider's representations, warranties, or obligations under this Agreement;",
        "b) Any infringement or misappropriation of any third party's Intellectual Property rights by the Services or Deliverables;",
        "c) Any Personal Data breach caused by the Service Provider's acts or omissions;",
        "d) Any claim by the Service Provider's employees, contractors, or subcontractors relating to employment, compensation, or working conditions;",
        "e) Any violation of applicable laws or regulations by the Service Provider in the performance of the Services;",
        "f) Any bodily injury, death, or property damage caused by the Service Provider's negligence or willful misconduct."
    ]
    for item in sp_indemnity:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, '10.2 Client Indemnification. The Client shall indemnify, defend, and hold harmless the Service Provider from and against any claims arising out of: (a) the Client\'s breach of this Agreement; (b) infringement of third-party rights by Client IP or Client-provided materials; or (c) the Client\'s violation of applicable laws.')

    add_para(doc, '10.3 Indemnification Procedure. The indemnified Party shall: (a) promptly notify the indemnifying Party in writing of any claim; (b) grant the indemnifying Party sole control of the defense and settlement; and (c) provide reasonable cooperation at the indemnifying Party\'s expense. The indemnifying Party shall not settle any claim that imposes obligations on the indemnified Party without prior written consent.')

    # Section 11: Insurance
    add_heading(doc, "11. Insurance", level=1)

    add_para(doc, '11.1 Required Coverage. The Service Provider shall maintain, at its own expense, the following insurance coverage throughout the Term and for a period of two (2) years following termination:')

    add_table(doc,
        ["Coverage Type", "Minimum Coverage", "Deductible (max)"],
        [
            ["Cyber Liability / Data Breach", "USD 10,000,000 per occurrence", "USD 250,000"],
            ["Professional Liability (E&O)", "USD 5,000,000 per occurrence", "USD 100,000"],
            ["Commercial General Liability", "USD 2,000,000 per occurrence", "USD 50,000"],
            ["Workers' Compensation", "As required by applicable law", "Statutory"],
            ["Employer's Liability", "USD 1,000,000 per occurrence", "USD 25,000"],
        ]
    )

    add_para(doc, '11.2 Evidence of Insurance. The Service Provider shall provide certificates of insurance to the Client within thirty (30) days of the Effective Date and annually thereafter. The Client shall be named as an additional insured on the Commercial General Liability and Cyber Liability policies.')

    add_para(doc, '11.3 Notice of Changes. The Service Provider shall provide at least thirty (30) days\' prior written notice to the Client of any material change, cancellation, or non-renewal of any required insurance coverage.')

    # Section 12: Termination
    add_heading(doc, "12. Termination", level=1)

    add_para(doc, '12.1 Termination for Convenience. Either Party may terminate this Agreement for convenience by providing one hundred and eighty (180) days\' prior written notice to the other Party. In the event of termination for convenience by the Client, the Client shall pay the Service Provider: (a) all fees for Services performed through the termination date; (b) reasonable wind-down costs; and (c) an early termination fee equal to six (6) months of the Base Fee if termination occurs during the first three (3) years of the Initial Term.')

    add_para(doc, '12.2 Termination for Cause by Client. The Client may terminate this Agreement immediately upon written notice if the Service Provider:')

    cause_items = [
        "a) Commits a material breach of this Agreement and fails to cure such breach within thirty (30) days of receiving written notice;",
        "b) Fails to meet Service Levels resulting in Service Credits exceeding twenty percent (20%) of the monthly Base Fee for three (3) consecutive months;",
        "c) Experiences a change of control without the Client's prior written consent;",
        "d) Becomes insolvent, files for bankruptcy, or has a receiver appointed;",
        "e) Is found to have engaged in fraud, corruption, or other criminal activity;",
        "f) Suffers a material data breach caused by the Service Provider's gross negligence or willful misconduct."
    ]
    for item in cause_items:
        doc.add_paragraph(item, style="List Bullet")

    add_para(doc, '12.3 Termination for Cause by Service Provider. The Service Provider may terminate this Agreement upon sixty (60) days\' written notice if the Client: (a) fails to pay undisputed invoices within ninety (90) days of the due date; or (b) commits a material breach that remains uncured for sixty (60) days after written notice.')

    add_para(doc, '12.4 Step-In Rights. In the event of a material service failure or the Service Provider\'s inability to perform, the Client shall have the right to step in and assume direct control of the affected Services, either through its own resources or a third-party provider, at the Service Provider\'s cost, until the Service Provider demonstrates its ability to resume performance.')

    add_para(doc, '12.5 Consequences of Termination. Upon termination or expiration of this Agreement: (a) the Service Provider shall immediately cease all Services except transition assistance; (b) all licenses granted to the Service Provider shall terminate; (c) the Service Provider shall return or destroy all Client IP and Confidential Information; and (d) all accrued rights and obligations shall survive.')

    add_para(doc, '12.6 Transition Assistance. The Service Provider shall provide transition assistance services for a period of up to twelve (12) months following the effective date of termination (the "Transition Period"). During the Transition Period, the Service Provider shall: (a) continue to provide the Services at the same Service Levels; (b) cooperate fully with the Client and any successor provider; (c) transfer all knowledge, documentation, and access credentials; (d) provide dedicated transition resources at no additional cost for the first six (6) months; and (e) make available transition resources at cost-plus-ten-percent (10%) for the remaining six (6) months.')

    # Section 13: Force Majeure
    add_heading(doc, "13. Force Majeure", level=1)

    add_para(doc, '13.1 Excuse of Performance. Neither Party shall be liable for any failure or delay in the performance of its obligations under this Agreement to the extent that such failure or delay is caused by a Force Majeure Event, provided that the affected Party: (a) gives prompt written notice to the other Party describing the Force Majeure Event and its expected duration; (b) uses commercially reasonable efforts to mitigate the effects of the Force Majeure Event; and (c) resumes performance as soon as reasonably practicable after the Force Majeure Event ceases.')

    add_para(doc, '13.2 Business Continuity. The Service Provider shall maintain and test a comprehensive business continuity and disaster recovery plan at least annually. The plan shall ensure recovery of critical services within: (a) four (4) hours for Severity 1 systems (Recovery Time Objective); and (b) one (1) hour of data loss tolerance (Recovery Point Objective). The Service Provider shall provide the Client with a copy of the plan and test results upon request.')

    add_para(doc, '13.3 Prolonged Force Majeure. If a Force Majeure Event continues for more than sixty (60) consecutive days, either Party may terminate this Agreement by providing thirty (30) days\' written notice to the other Party. In such event: (a) neither Party shall be liable for damages arising from the termination; (b) the Client shall pay for all Services performed up to the termination date; and (c) the Service Provider shall provide reasonable transition assistance for a period of ninety (90) days at no additional cost.')

    add_para(doc, '13.4 Pandemic Provisions. In the event of a pandemic or epidemic declared by the World Health Organization or relevant national health authority, the Service Provider shall: (a) activate its remote working protocols within forty-eight (48) hours; (b) maintain service delivery at no less than ninety percent (90%) of normal capacity within five (5) Business Days; and (c) provide weekly status reports on workforce availability and service impact.')

    # Section 14: Change Management
    add_heading(doc, "14. Change Management", level=1)

    add_para(doc, '14.1 Change Request Process. Either Party may request changes to the Services by submitting a formal Change Request. Each Change Request shall include: (a) description of the proposed change; (b) business justification; (c) impact assessment (scope, timeline, cost, resources, risk); (d) proposed implementation plan; and (e) any changes to Service Levels.')

    add_para(doc, '14.2 Evaluation and Approval. The Service Provider shall evaluate each Change Request and provide a written response within ten (10) Business Days, including: (a) technical feasibility assessment; (b) detailed cost estimate; (c) timeline for implementation; and (d) any impact on existing Services or Service Levels. No Change Request shall be implemented without the written approval of both Parties\' authorized representatives.')

    add_para(doc, '14.3 Emergency Changes. In the event of an emergency requiring immediate action to prevent harm to the Client\'s systems, data, or operations, the Service Provider may implement changes without prior formal approval, provided that: (a) the Service Provider notifies the Client within two (2) hours; and (b) a formal Change Request is submitted within twenty-four (24) hours for retrospective approval.')

    # Section 15: Compliance
    add_heading(doc, "15. Regulatory Compliance", level=1)

    add_para(doc, '15.1 General Compliance. The Service Provider shall comply with all applicable laws, regulations, and industry standards in the performance of the Services, including but not limited to:')

    add_table(doc,
        ["Standard / Regulation", "Requirement", "Certification Frequency"],
        [
            ["SOC 2 Type II", "Security, Availability, Confidentiality", "Annual"],
            ["ISO 27001:2022", "Information Security Management System", "Triennial (annual surveillance)"],
            ["ISO 27701:2019", "Privacy Information Management", "Triennial (annual surveillance)"],
            ["PCI DSS v4.0", "Payment Card Industry Data Security", "Annual"],
            ["HIPAA", "Health Insurance Portability and Accountability (if applicable)", "Annual assessment"],
            ["SOX", "Sarbanes-Oxley IT controls (if applicable)", "Annual"],
        ]
    )

    add_para(doc, '15.2 Anti-Corruption. The Service Provider represents and warrants that it has not and shall not, directly or indirectly, offer, pay, promise, or authorize the payment of any money or anything of value to any government official, political party, or candidate for the purpose of influencing any act or decision, in violation of the U.S. Foreign Corrupt Practices Act (FCPA), the UK Bribery Act 2010, or any other applicable anti-corruption law.')

    add_para(doc, '15.3 Export Controls. The Service Provider shall comply with all applicable export control laws and regulations, including the U.S. Export Administration Regulations (EAR) and International Traffic in Arms Regulations (ITAR). The Service Provider shall not export or re-export any technical data, software, or services provided under this Agreement to any country, entity, or person prohibited by applicable export control laws.')

    # Section 16: Dispute Resolution
    add_heading(doc, "16. Dispute Resolution", level=1)

    add_para(doc, '16.1 Escalation. Any dispute arising out of or in connection with this Agreement shall first be referred to the Parties\' respective project managers for resolution. If the dispute is not resolved within fifteen (15) Business Days, it shall be escalated to the Parties\' respective senior executives (VP level or above) for resolution within an additional fifteen (15) Business Days.')

    add_para(doc, '16.2 Mediation. If the dispute remains unresolved after the escalation process, the Parties shall attempt to resolve the dispute through mediation administered by the International Chamber of Commerce (ICC) in accordance with its Mediation Rules. The mediation shall take place in New York, New York.')

    add_para(doc, '16.3 Arbitration. If mediation fails to resolve the dispute within sixty (60) days, the dispute shall be finally resolved by binding arbitration administered by the ICC under its Rules of Arbitration. The arbitration shall be conducted by three (3) arbitrators, with each Party selecting one arbitrator and the two selected arbitrators selecting the third. The seat of arbitration shall be New York, New York. The language of the arbitration shall be English. The arbitral award shall be final and binding and may be enforced in any court of competent jurisdiction.')

    add_para(doc, '16.4 Injunctive Relief. Notwithstanding the foregoing, either Party may seek injunctive or other equitable relief from any court of competent jurisdiction to prevent irreparable harm, including but not limited to breaches of confidentiality, data protection, or Intellectual Property obligations.')

    # Section 17: Governing Law
    add_heading(doc, "17. Governing Law", level=1)

    add_para(doc, '17.1 Applicable Law. This Agreement shall be governed by and construed in accordance with the laws of the State of New York, United States, without regard to its conflict of laws principles.')

    add_para(doc, '17.2 Jurisdiction. Subject to Section 16 (Dispute Resolution), the Parties submit to the exclusive jurisdiction of the state and federal courts located in the Borough of Manhattan, New York, New York, for any legal proceedings arising out of this Agreement that are not subject to arbitration.')

    # Section 18: General Provisions
    add_heading(doc, "18. General Provisions", level=1)

    add_para(doc, '18.1 Assignment. Neither Party may assign or transfer this Agreement or any rights or obligations hereunder without the prior written consent of the other Party, except that either Party may assign this Agreement to an affiliate or in connection with a merger, acquisition, or sale of all or substantially all of its assets, provided that the assignee assumes all obligations under this Agreement.')

    add_para(doc, '18.2 Subcontracting. The Service Provider may subcontract portions of the Services to pre-approved subcontractors listed in Schedule 6, provided that: (a) the Service Provider remains fully responsible for the performance of the subcontracted Services; (b) all subcontractors are bound by obligations no less restrictive than those in this Agreement; and (c) the Client is notified of any new subcontractor at least thirty (30) days in advance.')

    add_para(doc, '18.3 Notices. All notices under this Agreement shall be in writing and delivered by: (a) personal delivery; (b) nationally recognized overnight courier; (c) registered or certified mail, return receipt requested; or (d) email with confirmed receipt. Notices shall be addressed to:')

    add_para(doc, 'If to the Client: Meridian Global Holdings, Inc., Attn: General Counsel, 500 Park Avenue, Suite 3200, New York, NY 10022, Email: legal@meridianglobal.com')
    add_para(doc, 'If to the Service Provider: TechServe Solutions Private Limited, Attn: Chief Legal Officer, Tower B, Cyber City, DLF Phase 2, Gurugram, Haryana 122002, India, Email: legal@techserve.in')

    add_para(doc, '18.4 Entire Agreement. This Agreement, together with all Schedules and Exhibits attached hereto, constitutes the entire agreement between the Parties with respect to the subject matter hereof and supersedes all prior negotiations, representations, warranties, commitments, offers, and agreements, whether written or oral.')

    add_para(doc, '18.5 Amendment. No amendment, modification, or waiver of any provision of this Agreement shall be effective unless in writing and signed by authorized representatives of both Parties.')

    add_para(doc, '18.6 Severability. If any provision of this Agreement is held to be invalid, illegal, or unenforceable, the remaining provisions shall continue in full force and effect.')

    add_para(doc, '18.7 Waiver. The failure of either Party to enforce any provision of this Agreement shall not constitute a waiver of such provision or the right to enforce it at a later time.')

    add_para(doc, '18.8 Counterparts. This Agreement may be executed in counterparts, each of which shall be deemed an original, and all of which together shall constitute one and the same instrument. Electronic signatures shall be deemed original signatures for all purposes.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    add_para(doc, "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.", bold=True)
    doc.add_paragraph()

    sig_table = doc.add_table(rows=5, cols=2)
    sig_table.style = "Table Grid"

    sig_table.rows[0].cells[0].text = "MERIDIAN GLOBAL HOLDINGS, INC."
    sig_table.rows[0].cells[1].text = "TECHSERVE SOLUTIONS PRIVATE LIMITED"
    for cell in sig_table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    sig_table.rows[1].cells[0].text = "Signature: _________________________"
    sig_table.rows[1].cells[1].text = "Signature: _________________________"

    sig_table.rows[2].cells[0].text = "Name: Jonathan R. Whitfield"
    sig_table.rows[2].cells[1].text = "Name: Rajesh K. Venkataraman"

    sig_table.rows[3].cells[0].text = "Title: Chief Information Officer"
    sig_table.rows[3].cells[1].text = "Title: Managing Director"

    sig_table.rows[4].cells[0].text = "Date: January 15, 2024"
    sig_table.rows[4].cells[1].text = "Date: January 15, 2024"

    return doc


if __name__ == "__main__":
    doc = build_contract()

    demo_path = os.path.join(os.path.dirname(__file__), "..", "demo", "samples", "complex_it_outsourcing.docx")
    demo_path = os.path.abspath(demo_path)
    os.makedirs(os.path.dirname(demo_path), exist_ok=True)
    doc.save(demo_path)
    print(f"Saved: {demo_path}")

    fixtures_path = os.path.join(os.path.dirname(__file__), "..", "tests", "fixtures", "complex_it_outsourcing.docx")
    fixtures_path = os.path.abspath(fixtures_path)
    os.makedirs(os.path.dirname(fixtures_path), exist_ok=True)
    doc.save(fixtures_path)
    print(f"Saved: {fixtures_path}")
